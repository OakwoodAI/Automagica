def activity(func):
    """Wrapper for Automagica activities
    """
    from functools import wraps
    import logging

    @wraps(func)
    def wrapper(*args, **kwargs):
        """Wrapper function
        """
        if func.__doc__:
            name = func.__doc__.split("\n")[0]
        else:
            name = func.__name__
        logging.info("Automagica (activity): {}".format(name))
        telemetry(func)
        return func(*args, **kwargs)

    return wrapper


def telemetry(func):
    """Automagica Activity Telemetry

    This allows us to collect information on the usage of 
    certain Automagica functionalities in order for us to keep improving 
    the software. If you would like to disable telemetry, make sure the 
    environment variable 'AUTOMAGICA_NO_TELEMETRY' is set. That way no
    information is being shared with us.
    """
    import requests
    from uuid import getnode
    import os
    import platform

    if not os.environ.get("AUTOMAGICA_NO_TELEMETRY") and not os.environ.get(
        "AUTOMAGICA_URL"
    ):
        if func.__doc__:
            name = func.__doc__.split("\n")[0]
        else:
            name = func.__name__

        data = {
            "activity": name,
            "machine_id": getnode(),
            "os": {
                "name": os.name,
                "platform": platform.system(),
                "release": platform.release(),
            },
        }

        try:
            r = requests.post(
                "https://telemetry.automagica.com/", json=data, timeout=1
            )
        except:
            pass


"""
Cryptography
Icon: las la-shield-alt
"""


@activity
def generate_random_key():
    """Random key

    Generate random Fernet key. Fernet guarantees that a message encrypted using it cannot be manipulated or read without the key. Fernet is an implementation of symmetric (also known as “secret key”) authenticated cryptography

    :return: Bytes-like object

        :Example:

    >>> generate_random_key()
    b'AYv6ZPVgnrUtHDbGZqAopRyAo9r0_UKrA2Rm3K_NjIo='

    Keywords
        random, key, fernet, hash, security, cryptography, password, secure

    Icon
        las la-key
    """
    import os
    from cryptography.fernet import Fernet
    key = Fernet.generate_key()

    return key


@activity
def encrypt_text_with_key(text, key):
    """Encrypt text 

    Encrypt text with (Fernet) key, 

    :parameter text: Text to be encrypted.
    :parameter path: Path where key is stored.

    :return: bytes-like object.

        :Example:

    >>> key = generate_random_key()
    >>> encrypt_text_with_key('Sample text', key)
    b'gAAAAABd8lpG8fNqcj5eXrPPHlx4KeCm-1TgX3jkyhStMfIlgGImIa-qaINZAj8XcxPcG8iu84iT56b_qAW9c5qpe7btUFhtxQ=='

    Keywords
        random, encryption, secure, security, hash, password, fernet, text

    Icon
        las la-lock
    """
    from cryptography.fernet import Fernet
    f = Fernet(key)

    return f.encrypt(text.encode('utf-8'))


@activity
def decrypt_text_with_key(encrypted_text, key):
    """Decrypt text

    Dexrypt bytes-like object to string with (Fernet) key

    :return: String

    :parameter encrypted_text: Text to be encrypted.
    :parameter path: Path where key is stored.

        :Example:

    >>> key = generate_random_key()
    >>> encrypted_text = encrypt_text_with_key('Sample text', key)
    >>> decrypt_text_with_key(encrypted_text, key)
    'Sample text'

    Keywords
        decrypt, random, unlock, un-lock hash, security, cryptography, password, secure, hash, text

    Icon
        las la-lock-open
    """
    from cryptography.fernet import Fernet
    f = Fernet(key)

    return f.decrypt(encrypted_text).decode("utf-8")


@activity
def encrypt_file_with_key(input_path, key, output_path=None):
    """Encrypt file 

    Encrypt file with (Fernet) key. Note that file will be unusable unless unlocked with the same key.

    :parameter input_file: Full path to file to be encrypted.
    :parameter output_file: Output path. Default is the same directory with "_encrypted" added to the name

    :return: Path to encrypted file

        :Example:

    >>> key = generate_random_key()
    >>> # Create a textfile to illustrate file encryption
    >>> textfile_path = make_textfile()
    >>> # Encrypt the textfile
    >>> encrypt_file_with_key(textfile_path, key=key)
    'C:\\Users\\<username>\\generated_textfile_encrypted.txt'

    Keywords
        encrypt, random, password, secure, secure file, lock

    Icon
        las la-lock
    """

    # Set path if not specified
    import os
    if not output_path:
        filepath = os.path.dirname(input_path)
        base = os.path.basename(input_path)
        filename = os.path.splitext(base)[0]
        extension = os.path.splitext(base)[1]
        output_path = os.path.join(
            filepath, filename + '_encrypted' + extension)

    from cryptography.fernet import Fernet
    with open(input_path, 'rb') as f:
        data = f.read()

    fernet = Fernet(key)
    encrypted = fernet.encrypt(data)

    with open(output_path, 'wb') as f:
        f.write(encrypted)

    return output_path


@activity
def decrypt_file_with_key(input_path, key, output_path=None):
    """Decrypt file

    Decrypts file with (Fernet) key

    :parameter input_file: Bytes-like file to be decrypted.
    :parameter output_file: Outputfile, make sure to give this the same extension as basefile before encryption. Default is the same directory with "_decrypted" added to the name 

    :return: Path to decrypted file

        :Example:

    >>> key = generate_random_key()
    >>> # Create a textfile to encrypt file
    >>> textfile_path = make_textfile()
    >>> # Encrypt the textfile
    >>> encrypted_textfile = encrypt_file_with_key(textfile_path, key=key)
    >>> # Decrypt the newly encrypted file
    >>> decrypt_file_with_key(encrypted_textfile, key=key)
    'C:\\Users\\<username>\\generated_textfile_encrypted_decrypted.txt'

    Keywords
        decrypt, random, password, secure, secure file, unlock

    Icon
        las la-lock-open
    """
    # Set path if not specified
    import os
    if not output_path:
        filepath = os.path.dirname(input_path)
        base = os.path.basename(input_path)
        filename = os.path.splitext(base)[0]
        extension = os.path.splitext(base)[1]
        output_path = os.path.join(
            filepath, filename + '_decrypted' + extension)

    from cryptography.fernet import Fernet
    with open(input_path, 'rb') as f:
        data = f.read()

    fernet = Fernet(key)
    decrypted = fernet.decrypt(data)

    with open(output_path, 'wb') as f:
        f.write(decrypted)

    return output_path


@activity
def generate_key_from_password(password, salt=None):
    """Key from password

    Generate key based on password and salt. If both password and salt are known the key can be regenerated.

    :parameter password: Passwords
    :parameter salt: Salt to generate key in combination with password. Default value is the hostname. Take in to account that hostname is necessary to generate key, e.g. when files are encrypted with salt 'A' and password 'B', both elements are necessary to decrypt files.

    :return: Bytes-like object

        :Example:

    >>> # Generate a key from password
    >>> key = generate_key_from_password(password='Sample password')
    b'7jGGF5w_xyI0CIZGCmLlnNyUvFpNvIUY08JCHopgAmm8='

    Keywords
        random, key, fernet, hash, security, cryptography, password, secure, salt

    Icon
        las la-lock
    """
    import base64
    from cryptography.hazmat.backends import default_backend
    from cryptography.hazmat.primitives import hashes
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    import socket

    # If no salt is set, use hostname as salt
    if not salt:
        salt = socket.gethostname().encode('utf-8')

    kdf = PBKDF2HMAC(algorithm=hashes.SHA256(), length=32,
                     salt=salt, iterations=500000, backend=default_backend())
    key = base64.urlsafe_b64encode(kdf.derive(password.encode('utf-8')))

    return key


@activity
def generate_hash_from_file(input_path, method='md5', buffer_size=65536):
    """Hash from file

    Generate hash from file 

    Can be used to create unique identifier for file validation or comparison.

    :parameter file: File to hash
    :parameter method: Method for hashing, choose between 'md5', 'sha256' and 'blake2b'. Note that different methods generate different hashes. Default method is 'md5'.
    :parameter buffer_size: Buffer size for reading file in chunks, default value is 64kb

    :return: Bytes-like object

        :Example:

    >>> # Generate a text file to illustrate hash
    >>> textfile_path = make_textfile()
    >>> # Get hash from textfile
    >>> generate_hash_from_file(textfile_path)
    '1ba249ca5931f3c85fe44d354c2f274d'

    Keywords
        hash, mdf5, sha256, blake2b, identifier, unique, hashing, fingerprint, comparison

    Icon
        las la-fingerprint

    """
    import sys
    import hashlib

    # Arbitrary buffer size. 64kb for compatibility with most systems
    buffer_size = 65536

    if method == 'md5':
        hash_list = hashlib.md5()
    if method == 'sha256':
        hash_list = hashlib.sha1()
    if method == 'blake2b':
        hash_list = hashlib.blake2b()

    with open(input_path, 'rb') as f:
        while True:
            data = f.read(buffer_size)
            if data:
                hash_list.update(data)
            else:
                return hash_list.hexdigest()


@activity
def generate_hash_from_text(text, method='md5'):
    """Hash from text

    Generate hash from text

    :parameter file: Text to hash
    :parameter method: Method for hashing, choose between 'md5', 'sha256' and 'blake2b'. Note that different methods generate different hashes. Default method is 'md5'.


        :Example:

    >>> # Generate a hast from text
    >>> generate_hash_from_text('Sample text')
    '1ba249ca5931f3c85fe44d354c2f274d'

    Keywords
        Hash, mdf5, sha256, blake2b, identifier, unique, hashing, fingerprint, text, comparison

    Icon
        las la-fingerprint

    """
    import sys
    import hashlib

    encoded_text = text.encode('utf-8')

    if method == 'md5':
        return hashlib.md5(encoded_text).hexdigest()
    if method == 'sha256':
        return hashlib.sha256(encoded_text).hexdigest()
    if method == 'blake2b':
        return hashlib.balke2b(encoded_text).hexdigest()


"""
Random
Icon: las la-dice-d6
"""


@activity
def generate_random_number(lower_limit=0, upper_limit=100, fractional=False):
    """Random number

    Random numbers can be integers (not a fractional number) or a float (fractional number).

    :parameter lower_limit: Lower limit for random number
    :parameter upper_limit: Upper limit for random number
    :parameter fractional: Setting this to True will generate fractional number. Default value is False and only generates whole numbers.

    :return: Random integer or float

        :Example:

    >>> # Generate a random number
    >>> generate_random_number()
    7

    Keywords:
        random number, random integer, dice, gamble, rng, random

    Icon
        las la-dice
    """
    import random
    if fractional:
        return random.uniform(lower_limit, upper_limit)
    else:
        return random.randrange(lower_limit, upper_limit, 1)


@activity
def generate_random_boolean():
    """Random boolean

    Generates a random boolean (True or False)

    :return: Boolean

        :Example:

    >>> # Generate a random boolean
    >>> generate_random_boolean()
    True

    Keywords:
        random, dice, gamble, rng, coin, coinflip, heads, tails

    Icon
        las la-coins
    """
    import random
    return bool(random.getrandbits(1))


@activity
def generate_random_name(locale=None):
    """Random name

    Generates a random name. Adding a locale adds a more common name in the specified locale. Provides first name and last name.

    :parameter locale: Add a locale to generate popular name for selected locale.

        -   ar_EG - Arabic (Egypt)
        -   ar_PS - Arabic (Palestine)
        -   ar_SA - Arabic (Saudi Arabia)
        -   bg_BG - Bulgarian
        -   bs_BA - Bosnian
        -   cs_CZ - Czech
        -   de_DE - German
        -   dk_DK - Danish
        -   el_GR - Greek
        -   en_AU - English (Australia)
        -   en_CA - English (Canada)
        -   en_GB - English (Great Britain)
        -   en_NZ - English (New Zealand)
        -   en_US - English (United States)
        -   es_ES - Spanish (Spain)
        -   es_MX - Spanish (Mexico)
        -   et_EE - Estonian
        -   fa_IR - Persian (Iran)
        -   fi_FI - Finnish
        -   fr_FR - French
        -   hi_IN - Hindi
        -   hr_HR - Croatian
        -   hu_HU - Hungarian
        -   hy_AM - Armenian
        -   it_IT - Italian
        -   ja_JP - Japanese
        -   ka_GE - Georgian (Georgia)
        -   ko_KR - Korean
        -   lt_LT - Lithuanian
        -   lv_LV - Latvian
        -   ne_NP - Nepali
        -   nl_NL - Dutch (Netherlands)
        -   no_NO - Norwegian
        -   pl_PL - Polish
        -   pt_BR - Portuguese (Brazil)
        -   pt_PT - Portuguese (Portugal)
        -   ro_RO - Romanian
        -   ru_RU - Russian
        -   sl_SI - Slovene
        -   sv_SE - Swedish
        -   tr_TR - Turkish
        -   uk_UA - Ukrainian
        -   zh_CN - Chinese (China)
        -   zh_TW - Chinese (Taiwan)

    :return: Name as string

        :Example:

    >>> # Generate a random name
    >>> generate_random_name()
    'Michelle Murphy'

    Keywords
        random, dummy name, name, name generater, fake person, fake, person, surname, lastname, fake name generator

    Icon
        las la-user-tag
    """
    from faker import Faker
    if locale:
        seed = Faker(locale)
    else:
        seed = Faker()
    return seed.name()


def generate_random_sentence(locale=None):
    """Random sentence

    Generates a random sentence. Specifying locale changes language and content based on locale.

    :parameter locale: Add a locale to generate popular name for selected locale.

        -   ar_EG - Arabic (Egypt)
        -   ar_PS - Arabic (Palestine)
        -   ar_SA - Arabic (Saudi Arabia)
        -   bg_BG - Bulgarian
        -   bs_BA - Bosnian
        -   cs_CZ - Czech
        -   de_DE - German
        -   dk_DK - Danish
        -   el_GR - Greek
        -   en_AU - English (Australia)
        -   en_CA - English (Canada)
        -   en_GB - English (Great Britain)
        -   en_NZ - English (New Zealand)
        -   en_US - English (United States)
        -   es_ES - Spanish (Spain)
        -   es_MX - Spanish (Mexico)
        -   et_EE - Estonian
        -   fa_IR - Persian (Iran)
        -   fi_FI - Finnish
        -   fr_FR - French
        -   hi_IN - Hindi
        -   hr_HR - Croatian
        -   hu_HU - Hungarian
        -   hy_AM - Armenian
        -   it_IT - Italian
        -   ja_JP - Japanese
        -   ka_GE - Georgian (Georgia)
        -   ko_KR - Korean
        -   lt_LT - Lithuanian
        -   lv_LV - Latvian
        -   ne_NP - Nepali
        -   nl_NL - Dutch (Netherlands)
        -   no_NO - Norwegian
        -   pl_PL - Polish
        -   pt_BR - Portuguese (Brazil)
        -   pt_PT - Portuguese (Portugal)
        -   ro_RO - Romanian
        -   ru_RU - Russian
        -   sl_SI - Slovene
        -   sv_SE - Swedish
        -   tr_TR - Turkish
        -   uk_UA - Ukrainian
        -   zh_CN - Chinese (China)
        -   zh_TW - Chinese (Taiwan)

    :return: Random sentence as string

        :Example:

    >>> # Generate a random sentence
    >>> generate_random_sentence()
    'The age of automation is going to be the age of do-it-yourself'

    Keywords
        random, sentence, lorem ipsum, text generater, filler, place holder, noise, random text

    Icon
        las la-comment
    """
    from faker import Faker
    if locale:
        seed = Faker(locale)
    else:
        seed = Faker()
    return seed.sentence()


@activity
def generate_random_address(locale=None):
    """Random address

    Generates a random address. Specifying locale changes random locations and streetnames based on locale.

    :parameter locale: Add a locale to generate popular name for selected locale.

        -   ar_EG - Arabic (Egypt)
        -   ar_PS - Arabic (Palestine)
        -   ar_SA - Arabic (Saudi Arabia)
        -   bg_BG - Bulgarian
        -   bs_BA - Bosnian
        -   cs_CZ - Czech
        -   de_DE - German
        -   dk_DK - Danish
        -   el_GR - Greek
        -   en_AU - English (Australia)
        -   en_CA - English (Canada)
        -   en_GB - English (Great Britain)
        -   en_NZ - English (New Zealand)
        -   en_US - English (United States)
        -   es_ES - Spanish (Spain)
        -   es_MX - Spanish (Mexico)
        -   et_EE - Estonian
        -   fa_IR - Persian (Iran)
        -   fi_FI - Finnish
        -   fr_FR - French
        -   hi_IN - Hindi
        -   hr_HR - Croatian
        -   hu_HU - Hungarian
        -   hy_AM - Armenian
        -   it_IT - Italian
        -   ja_JP - Japanese
        -   ka_GE - Georgian (Georgia)
        -   ko_KR - Korean
        -   lt_LT - Lithuanian
        -   lv_LV - Latvian
        -   ne_NP - Nepali
        -   nl_NL - Dutch (Netherlands)
        -   no_NO - Norwegian
        -   pl_PL - Polish
        -   pt_BR - Portuguese (Brazil)
        -   pt_PT - Portuguese (Portugal)
        -   ro_RO - Romanian
        -   ru_RU - Russian
        -   sl_SI - Slovene
        -   sv_SE - Swedish
        -   tr_TR - Turkish
        -   uk_UA - Ukrainian
        -   zh_CN - Chinese (China)
        -   zh_TW - Chinese (Taiwan)

    :return: Random address as string

        :Example:

    >>> # Generate a random address
    >>> generate_random_address()
    '5639 Cynthia Bridge Suite 610
    'Port Nancy, GA 95894'

    Keywords
        random, address, random address, fake person , fake address, fake person generator

    Icon
        las la-map
    """
    from faker import Faker
    if locale:
        seed = Faker(locale)
    else:
        seed = Faker()
    return seed.address()


@activity
def generate_random_beep(max_duration=2000, max_frequency=5000):
    """Random beep

    Generates a random beep, only works on Windows

    :parameter max_duration: Maximum random duration in miliseconds. Default value is 2 miliseconds
    :parameter max_frequency: Maximum random frequency in Hz. Default value is 5000 Hz.

    :return: Sound

        :Example: 

    >>> # Generate a random beep
    >>> generate_random_beep()

    Keywords
        beep, sound, random, noise, alert, notification

    Icon
        las la-volume-up
    """
    import winsound
    import random
    frequency = random.randrange(5000)
    duration = random.randrange(2000)
    winsound.Beep(frequency, duration)


@activity
def generate_random_date(format='%m/%d/%Y %I:%M', days_in_past=1000):
    """Random date

    Generates a random date.

    -   %a	Abbreviated weekday name.	 
    -   %A	Full weekday name.	 
    -   %b	Abbreviated month name.	 
    -   %B	Full month name.	 
    -   %c	Predefined date and time representation.	 
    -   %d	Day of the month as a decimal number [01,31].	 
    -   %H	Hour (24-hour clock) as a decimal number [00,23].	 
    -   %I	Hour (12-hour clock) as a decimal number [01,12].	 
    -   %j	Day of the year as a decimal number [001,366].	 
    -   %m	Month as a decimal number [01,12].	 
    -   %M	Minute as a decimal number [00,59].	 
    -   %p	AM or PM.
    -   %S	Second as a decimal number [00,61].	
    -   %U	Week number of the year (Sunday as the first day of the week) as a decimal number [00,53]. All days in a new year preceding the first Sunday are considered to be in week 0.	
    -   %w	Weekday as a decimal number [0(Sunday),6].	 
    -   %W	Week number of the year (Monday as the first day of the week) as a decimal number [00,53]. All days in a new year preceding the first Monday are considered to be in week 0.	
    -   %x	Predefined date representation.	 
    -   %X	Predefined time representation.	 
    -   %y	Year without century as a decimal number [00,99].	 
    -   %Y	Year with century as a decimal number.
    -   %Z	Time zone name (no characters if no time zone exists).

    :parameter days_in_past: Days in the past for which oldest random date is generated, default is 1000 days
    :parameter format: Formatting of the dates, replace with 'None' to get raw datetime format. e.g. format='Current month is %B' generates 'Current month is Januari' and format='%m/%d/%Y %I:%M' generates format 01/01/1900 00:00. 

    :return: Random date as string

        :Example: 

    >>> # Generate a random date
    >>> generate_random_date()
    01/01/2020 13:37'

    Keywords
        random, date, datetime, random date, fake date , calendar

    Icon
        las la-calendar
    """

    import random
    import datetime

    latest = datetime.datetime.now()
    earliest = latest - datetime.timedelta(days=days_in_past)
    delta_seconds = (latest - earliest).total_seconds()

    random_date = earliest + \
        datetime.timedelta(seconds=random.randrange(delta_seconds))

    if format:
        return random_date.strftime(format)
    else:
        return random_date


@activity
def generate_unique_identifier():
    """Generate unique identifier

    Generates a random UUID4 (universally unique identifier). While the probability that a UUID will be duplicated is not zero, it is close enough to zero to be negligible.

    :return: Identifier as string

        :Example:

    >>> # Generate unique identifier
    >>> generate_unique_identifier()
    'd72fd7ea-d682-4f78-8ca1-0ed34142a992'

    Keywords
        unique, identifier, primary key, random

    Icon
        las la-random
    """
    from uuid import uuid4

    return str(uuid4())


"""
User Input
Icon: lab la-wpforms
"""


@activity
def ask_user_input(title="Title", label="Input", password=False):
    """Ask user for input

    Prompt the user for an input with a pop-up window. 

    :parameter title: Title for the pop-up window
    :parameter message: The message to be shown to the user

    :return: Inputted text as string

        :Example:

    >>> # Make a window pop-up ask for user input
    >>> ask_user_input()
    >>> # Type in text and press 'submit', e.g. 'Sample text'
    'Sample text'

    Keywords
        user input, pop-up, interaction, popup, window, feedback, screen, ad-hoc, attended

    Icon
        las la-window-maximize
    """
    import PySimpleGUI as sg

    sg.ChangeLookAndFeel("SystemDefault")

    text = sg.Text(label, background_color="#2196F3", text_color="white")

    input_field = sg.InputText(justification="center", focus=True)

    if password:
        input_field = sg.InputText(
            password_char="*", justification="center", focus=True
        )

    submit_button = sg.Submit(button_color=("white", "#0069C0"))

    layout = [[text], [input_field], [submit_button]]

    window = sg.Window(
        title,
        layout,
        icon="icon.ico",
        no_titlebar=True,
        background_color="#2196F3",
        element_justification="center",
        use_default_focus=False,
        keep_on_top=True
    )
    _, values = window.Read()
    window.Close()
    value = values[0]

    return value


@activity
def ask_user_password(label="Password"):
    """Ask user for password

    Prompt the user for a password. The password will be masked on screen while entering.

    :parameter title: Title for the pop-up window
    :parameter message: The message to be shown to the user

    :return: Inputted password as string

        :Example:

    >>> # Make a window pop-up ask for user password
    >>> ask_user_password()
    >>> # Type in password and press 'submit', e.g. 'Sample password'
    'Sample password'

    Keywords
        user input, pop-up, interaction, interactive, credential, popup, window, feedback, password, screen, login, attended

    Icon
        lar la-window-maximize
    """
    return ask_user_input(title="Password", label=label, password=True)


@activity
def ask_credentials(title="Credentials required", dialogue_text_username="Username:", dialogue_text_password="Password:"):
    """Ask user for credentials

    Prompt a popup which asks user for username and password and returns in plain text. Password will be masked.

    :parameter title: Title for the popup
    :parameter dialogue_text: Dialogue text for username
    :parameter dialogue_text: Dialogue text for password

    :return: Typle with nputted username and password as strings

        :Example:

    >>> # Make a window pop-up ask user credentials
    >>> ask_credentials()
    >>> # Type in Username and Password 'submit', e.g. 'Sample username' and 'Sample password'
    ('Sample username', 'Sample password')

    Keywords
        user input, credentials, interactive, pop-up, interaction, popup, window, feedback, password, screen, login, attended

    Icon
        las la-window-maximize
    """
    import PySimpleGUI as sg

    sg.ChangeLookAndFeel("SystemDefault")

    layout = [
        [
            sg.Text(
                dialogue_text_username, background_color="#2196F3", text_color="white"
            )
        ],
        [sg.InputText(justification="center", focus=True)],
        [
            sg.Text(
                dialogue_text_password, background_color="#2196F3", text_color="white"
            )
        ],
        [sg.InputText(password_char="*", justification="center")],
        [sg.Submit(button_color=("white", "#0069C0"))],
    ]

    window = sg.Window(
        title,
        layout,
        icon="icon.ico",
        no_titlebar=True,
        background_color="#2196F3",
        element_justification="center",
        use_default_focus=False,
        keep_on_top=True
    )
    _, values = window.Read()

    window.Close()
    username = values[0]
    password = values[1]

    return username, password

@activity
def display_message_box(title="Title", message="Example message"):
    """Shows message box

    A pop-up message with title and message. 

    :parameter title: Title for the pop-up window
    :parameter message: The message to be shown to the user

    :return: True if user presses 'OK'

        :Example:

    >>> # Show pop-up with message
    >>> display_message_box()
    >>> # Wait till user presses 'OK'
    True

    Keywords
        message box, warning, info, popup, window, feedback, screen, attended

    Icon
        las la-window-maximize
    """
    import PySimpleGUI as sg

    sg.ChangeLookAndFeel("SystemDefault")

    text = sg.Text(message, background_color="#2196F3", text_color="white")

    ok_button = sg.OK(button_color=("white", "#0069C0"))

    layout = [[text], [ok_button]]

    window = sg.Window(
        title,
        layout,
        icon="icon.ico",
        no_titlebar=True,
        background_color="#2196F3",
        element_justification="center",
        use_default_focus=False,
        keep_on_top=True
    )
    _, values = window.Read()
    window.Close()

    return True


@activity
def display_osd_message(message='Example message', seconds=5):
    """Display overlay message

    Display custom OSD (on-screen display) message. Can be used to display a message for a limited amount of time. Can be used for illustration, debugging or as OSD.

    :parameter message: Message to be displayed

        :Example:

    >>> # Display overlay message
    >>> display_osd_message()

    Keywords
        message box, osd, overlay, info warning, info, popup, window, feedback, screen, login, attended

    Icon
        las la-tv
    """
    if "DISABLE_AUTOMAGICA_OSD" in globals():
        return

    from threading import Thread

    def load_osd():
        import tkinter
        import win32con
        import pywintypes
        import win32api

        screen_width = win32api.GetSystemMetrics(0)
        screen_height = win32api.GetSystemMetrics(1)

        root = tkinter.Tk()
        label = tkinter.Label(
            text=message, font=("Helvetica", "30"), fg="white", bg="black", borderwidth=10
        )
        label.master.overrideredirect(True)
        label.config(anchor=tkinter.CENTER)
        label.master.geometry(
            "+{}+{}".format(int(screen_width / 2), int(screen_height - 250))
        )
        label.master.lift()
        label.master.wm_attributes("-topmost", True)
        label.master.wm_attributes("-disabled", True)
        label.master.wm_attributes("-transparentcolor", "black")

        hWindow = pywintypes.HANDLE(int(label.master.frame(), 16))

        exStyle = (
            win32con.WS_EX_COMPOSITED
            | win32con.WS_EX_LAYERED
            | win32con.WS_EX_NOACTIVATE
            | win32con.WS_EX_TOPMOST
            | win32con.WS_EX_TRANSPARENT
        )
        win32api.SetWindowLong(hWindow, win32con.GWL_EXSTYLE, exStyle)

        label.after(seconds * 1000, lambda: root.destroy())
        label.pack()
        label.mainloop()

    t = Thread(target=load_osd)
    try:
        t.start()
    except:
        pass
    finally:
        try:
            t.kill()
        except:
            pass


"""
Browser
Icon: lab la-chrome
"""

import selenium.webdriver


class Chrome(selenium.webdriver.Chrome):
    @activity
    def __init__(self, load_images=True, headless=False):
        """Opens Chrome Browser

        Open the Chrome Browser with the Selenium webdriver. Canb be used to automate manipulations in the browser.
        Different elements can be found as:

        -   Xpath: e.g. browser.find_element_by_xpath() or browser.xpath()
        One can easily find an xpath by right clicking an element -> inspect. Look for the element in the menu and right click -> copy -> xpath
        find_element_by_id
        -   Name: find_element_by_name
        -   Link text: find_element_by_link_text
        -   Partial link text: find_element_by_partial_link_text
        -   Tag name: find_element_by_tag_name
        -   Class name: find_element_by_class_name
        -   Css selector: find_element_by_css_selector

        Elements can be manipulated by:

        - Clicking: e.g. element.click()
        - Typing: e.g. element.send_keys()

        :parameter load_images: Do not load images (bool). This could speed up loading pages
        :parameter headless: Run headless, this means running without a visible window (bool)

        return: wWbdriver: Selenium Webdriver

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://automagica.com')
        >>> # Close browser
        >>> browser.quit()

        Keywords
            chrome, browsing, browser, internet, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            lab la-chrome

        """
        import platform
        import os

        # Check what OS we are on
        if platform.system() == "Linux":
            chromedriver_path = "\\bin\\webdriver\\linux64\\chromedriver"
        elif platform.system() == "Windows":
            chromedriver_path = "\\bin\\win32\\chromedriver.exe"
        else:
            chromedriver_path = "\\bin\\mac64\\chromedriver.exe"

        chrome_options = selenium.webdriver.ChromeOptions()

        if headless:
            chrome_options.add_argument("--headless")

        if not load_images:
            prefs = {"profile.managed_default_content_settings.images": 2}
            chrome_options.add_experimental_option("prefs", prefs)

        selenium.webdriver.Chrome.__init__(self, os.path.abspath(
            __file__).replace('activities.py', '') + chromedriver_path)

    @activity
    def save_all_images(self, output_path=None):
        """Save all images

        Save all images on current page in the Browser

        :parameter output_path: Path where images can be saved. Default value is home directory.

        :return: List with paths to images

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://www.nytimes.com/')
        >>> # Save all images
        >>> browser.save_all_images()
        >>> browser.quit()
        ['C:\\Users\\<username>\\image1.png', 'C:\\Users\\<username>\\image2.jpg', 'C:\\Users\\<username>\\image4.gif']

        Keywords
            image scraping, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-images

        """
        import requests
        import os
        from urllib.parse import urlparse

        if not output_path:
            output_path = os.path.expanduser("~")

        paths = []

        images = self.find_elements_by_tag_name('img')

        for image in images:
            url = image.get_attribute('src')
            a = urlparse(url)
            filename = os.path.basename(a.path)

            with open(os.path.join(output_path, filename), 'wb') as f:
                try:
                    r = requests.get(url)
                    f.write(r.content)
                    paths.append(os.path.join(output_path, filename))
                except:
                    pass

        return paths

    @activity
    def find_elements_by_text(self, text):
        """Find elements by text

        Find all elements by their text. Text does not need to match exactly, part of text is enough.

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://nytimes.com')
        >>> # Find elements by text
        >>> browser.find_elements_by_text('world')
        [webelement1, webelement2 , .. ]

        Keywords
            element, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-align-center

        """
        return self.find_elements_by_xpath("//*[contains(text(), '" + text.lower() + "')] | //*[@value='" + text.lower() + "']")

    @activity
    def find_element_by_text(self, text):
        """Find element by text

        Find one element by text. Text does not need to match exactly, part of text is enough.

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://nytimes.com')
        >>> # Find elements by text
        >>> browser.find_element_by_text('world')
        webelement

        Keywords
            element, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-align-center

        """
        return self.find_element_by_xpath("//*[contains(text(), '" + text.lower() + "')] | //*[@value='" + text.lower() + "']")

    @activity
    def find_all_links(self):
        """Find all links

        Find all links on a webpage in the browser

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://nytimes.com')
        >>> # Find elements by text
        >>> browser.find_all_links()
        [webelement1, webelement2 , .. ]

        Keywords
            random, element, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-window-restore
        """
        return self.find_elements_by_xpath("//a[@href]")

    @activity
    def highlight(self, element):
        """Highlight element

        Highlight elements in yellow in the browser

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://wikipedia.org')
        >>> # Find all links
        >>> links = browser.find_all_links()
        >>> # Select first link to highlight for illustration
        >>> first_link = links[0]
        >>> # Highlight first link
        >>> browser.highlight(first_link)

        Keywords
            element, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-highlighter

        """
        driver = element._parent

        def apply_style(s):
            driver.execute_script("arguments[0].setAttribute('style', arguments[1]);",
                                  element, s)
        original_style = element.get_attribute('style')

        apply_style("background: yellow; border: 2px solid red;")

    @activity
    def exit(self):
        """Exit the browser

        Quit the browser by exiting gracefully. One can also use the native 'quit' function, e.g. 'browser.quit()'

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://automagica.com')
        >>> # Close browser
        >>> browser.exit()


        Keywords
            quit, exit, close, element, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-window-close
        """
        self.quit()

    @activity
    def find_all_xpaths(self, element):
        """Find all XPaths

        Find all elements with specified xpath on a webpage in the the browser. Can also use native 'find_elements_by_xpath' function e.g. browser.find_elements_by_xpath()
        You can easily

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://wikipedia.org')
        >>> # Find element by xpath
        >>> browser.find_xpaths('//*[@id=\'js-link-box-en\']')
        [webelement1, webelement2 , .. ]

        Keywords
            random, element, xpath, xml, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-times

        """
        return self.find_elements_by_xpath(element)

    @activity
    def find_xpath(self, element):
        """Find XPath in browser

        Find all element with specified xpath on a webpage in the the browser. Can also use native 'find_elements_by_xpath' function e.g. browser.find_element_by_xpath()

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://wikipedia.org')
        >>> # Find element by xpath
        >>> elements = browser.find_xpath('//*[@id=\'js-link-box-en\']')
        >>> # We can now use this element, for example to click on
        >>> element.click()

        Keywords
            random, xpath, element, xml element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-times
        """
        return self.find_element_by_xpath(element)


"""
Credential Management
Icon: las la-key
"""


@activity
def set_credential(username=None, password=None, system="Automagica"):
    """Set credential

    Add a credential which stores credentials locally and securely. All parameters should be Unicode text. 

    :parameter username: Username for which credential will be added.
    :parameter password: Password to add
    :parameter system: Name of the system for which credentials are stored. Extra safety measure and method for keeping passwords for similar usernames on different applications a part. Highly recommended to change default value.

    :return: Stores credentials locally

        :Example:

    >>> set_credential('SampleUsername', 'SamplePassword')

    Keywords
        credential, login, password, username, store, vault, secure, credentials, store, log in, encrypt

    Icon
        las la-key

    """
    import keyring

    keyring.set_password(system, username, password)


@activity
def delete_credential(username=None, password=None, system="Automagica"):
    """Delete credential

    Delete a locally stored credential. All parameters should be Unicode text. 

    :parameter username: Username for which credential (username + password) will be deleted.
    :parameter system: Name of the system for which password will be deleted. 


        :Example:

    >>> set_credential('SampleUsername', 'SamplePassword')
    >>> delete_credential('SampleUsername', 'SamplePassword')

    Keywords
        credential, delete, login, password, username, store, vault, secure, credentials, store, log in, encrypt

    Icon
        las la-key

    """
    import keyring

    keyring.delete_password(system, username)


@activity
def get_credential(username=None, system="Automagica"):
    """Get credential

    Get a locally stored redential. All parameters should be Unicode text. 

    :parameter username: Username to get password for.
    :parameter system: Name of the system for which credentials are retreived.

    :return: Stored credential as string

        :Example:

    >>> set_credential('SampleUsername', 'SamplePassword')
    >>> get_credential('SampleUsername')
    'SamplePassword'

    Keywords
        credential, get, delete, login, password, username, store, vault, secure, credentials, store, log in, encrypt

    Icon
        las la-key

    """
    import keyring

    return keyring.get_password(system, username)


"""
FTP
Icon: las la-key
"""


class FTP:
    @activity
    def __init__(self, server, username, password):
        """Create FTP connection

        Can be used to automate activites for FTP

        :parameter server: Name of the server
        :parameter username: Username 
        :parameter password: Password

            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Take caution uploading and downloading from this server as it is public
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')

        Keywords
            FTP, file transfer protocol, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-folder-open

        """
        import ftplib

        self.connection = ftplib.FTP(server)
        self.connection.login(username, password)

    @activity
    def download_file(self, input_path, output_path=None):
        """Download file
        
        Downloads a file from FTP server. Connection needs to be established first.

        :parameter input_path: Path to the file on the FPT server to download
        :parameter output_path: Destination path for downloaded files. Default is the same directory with "_downloaded" added to the name

        :return: Path to output file as string 

            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Take caution uploading and downloading from this server as it is public
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')
        >>> # Download Rebex public file 'readme.txt'
        >>> ftp.download_file('readme.txt')
        'C:\\Users\\<username>\\readme_downloaded.txt'

        Keywords
            FTP, file transfer protocol, download, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-download

        """
        # Set path if not specified
        if not output_path:
            import os
            filepath = os.path.expanduser("~")
            base = os.path.basename(input_path)
            filename = os.path.splitext(base)[0]
            extension = os.path.splitext(base)[1]
            output_path = os.path.join(
                filepath, filename + '_downloaded' + extension)

        self.connection.retrbinary(
            "RETR " + input_path, open(output_path, "wb").write)

        return output_path

    @activity
    def upload_file(self, input_path, output_path=None):
        """Upload file
        
        Upload file to FTP server

        :parameter from_path: Path file that will be uploaded
        :parameter to_path: Destination path to upload. 

        :return: Patht to uploaded file as string

            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Take caution uploading and downloading from this server as it is public
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')
        >>> # Create a .txt file for illustration
        >>> textfile = make_textfile()
        >>> # Upload file to FTP test server
        >>> # Not that this might result in a persmission error for public FPT's
        >>> ftp.upload_file(textfile)

        Keywords
            FTP, upload, fptfile transfer protocol, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-upload
        """
        # Set to user home if no path specified
        if not output_path:
            output_path = "/"

        self.connection.retrbinary(
            "RETR " + input_path, open(output_path, "wb").write)

    @activity
    def enumerate_files(self, path="/"):
        """List FTP files

        Generate a list of all the files in the FTP directory

        :parameter path: Path to list files from. Default is the main directory

        :return: Prints list of all files and directories

            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Take caution uploading and downloading from this server as it is public
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')
        >>> # Show all files in main directory
        >>> ftp.enumerate_files()
        10-27-15  03:46PM       <DIR>          pub
        04-08-14  03:09PM                  403 readme.txt
        '226 Transfer complete.'

        Keywords
            FTP, list, upload, fptfile transfer protocol, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-list-ol
        """
        self.connection.cwd(path)
        lines = self.connection.retrlines("LIST")
        return lines

    @activity
    def directory_exists(self, path="/"):
        """Check FTP directory
        
        Check if FTP directory exists

        :parameter path: Path to check on existence. Default is main directory

        :return: Boolean

            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Take caution uploading and downloading from this server as it is public
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')
        >>> # Check if 'pub' folder exists in main directory
        >>> ftp.directory_exists('\\pub')
        True

        Keywords
            FTP, list, upload, fptfile transfer protocol, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-list-ol

        """
        try:
            self.connection.cwd(path)
            return True
        except:
            return False

    @activity
    def create_directory(self, directory_name, path="/"):
        """Create FTP directory

        Create a FTP directory. Note that sufficient permissions are present

        :parameter directory_name: Name of the new directory, should be a string e.g. 'my_directory'
        :parameter path: Path to parent directory where to make new directory. Default is main directory

        :return: Boolean if creation was succesful (True) or failed (False)
            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Trying to create a directory will most likely fail due to permission
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')
        >>> # Create directory
        >>> ftp.create_directory('brand_new_directory')      
        False

        Keywords
            FTP, create, create folder, new, new folder, fptfile transfer protocol, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-folder-plus

        """
        try:
            self.connection.cwd(path)
            try:
                self.connection.mkd(directory_name)
                return True
            except Exception as e:
                if not e.args[0].startswith("550"):  # Exists already
                    raise
        except:
            return False


"""
Keyboard
Icon: las la-keyboard
"""


def easy_key_translation(key):
    """Activity supporting key translations
    """

    if not key:
        return ''

    key_translation = {'backspace': '{BACKSPACE}',
                       'break': '{BREAK}',
                       'capslock': '{CAPSLOCK}',
                       'delete': '{DELETE}',
                       'alt': '%',
                       'ctrl': '^',
                       'shift': '+',
                       'downarrow': '{DOWN}',
                       'end': '{END}',
                       'enter': '{ENTER}',
                       'escape': '{ESC}',
                       'help': '{HELP}',
                       'home': '{HOME}',
                       'insert': '{INSERT}',
                       'win': '^{Esc}',
                       'leftarrow': '{LEFT}',
                       'numlock': '{NUMLOCK}',
                       'pagedown': '{PGDN}',
                       'pageup': '{PGUP}',
                       'printscreen': '{PRTSC}',
                       'rightarrow': '{RIGHT}',
                       'scrolllock': '{SCROLLLOCK}',
                       'tab': '{TAB}',
                       'uparrow': '{UP}',
                       'f1': '{F1}',
                       'f2': '{F2}',
                       'f3': '{F3}',
                       'f4': '{F4}',
                       'f5': '{F5}',
                       'f6': '{F6}',
                       'f7': '{F7}',
                       'f8': '{F8}',
                       'f9': '{F9}',
                       'f10': '{F10}',
                       'f11': '{F11}',
                       'f12': '{F12}',
                       'f13': '{F13}',
                       'f14': '{F14}',
                       'f15': '{F15}',
                       'f16': '{F16}',
                       '+': '{+}',
                       '^': '{^}',
                       '%': '{%}',
                       '~': '{~}',
                       '(': '{(}',
                       ')': '{)}',
                       '[': '{[}',
                       ']': '{]}',
                       '{': '{{}',
                       '}': '{}}',
                       }

    if key_translation.get(key):
        return key_translation.get(key)

    return key


@activity
def press_key(key=None):
    """Press key

    Press and release an entered key. Make sure your keyboard is on US layout (standard QWERTY). 
    If you are using this on Mac Os you might need to grant acces to your terminal application. (Security Preferences > Security & Privacy > Privacy > Accessibility)

    Supported keys:
        ' ', '!', '"', '#', '$', '%', '&', "'", '(', ,')', '*', '+', ',', '-', '.', '/', '0', '1', '2', '3', '4', '5', '6', '7', '8', '9', ':', ';', '<','=', '>', '?', '@', '[', '\\', ']', '^', '_', '`', 'a', 'b', 'c', 'd', 'e','f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z', '{', '|', '}', '~', 'alt', 'backspace',  'ctrl', 'delete' 'downarrow', 'rightarrow', 'leftarrow', 'uparrow', 'enter', 'escape', 'f1', 'f2', f3', 'f4', 'f5', 'f6', 'f7', 'f8',  'f9', 'f10', 'f11', 'f12', 'f13', 'f14', 'f15', 'f16', 'home', 'insert', 'pagedown', 'pageup', 'help', 'printscreen', 'space', 'scrollock', 'tab', shift, 'win'

    :parameter key: Key to press

    :return: Keypress

        :Example:

    >>> # Open notepad to illustrate typing
    >>> run('notepad.exe')
    >>> # Press some keys
    >>> press_key('a')
    >>> press_key('enter')
    >>> press_key('b')
    >>> press_key('enter')
    >>> press_key('c')

    Keywords
        keyboard, typing, type, key, keystroke, hotkey, press, press key

    Icon
        las la-keyboard

    """
    import platform

    # Check if system is not running Windows
    if platform.system() != "Windows":
        from pyautogui import press

        if key:
            return press(key)

    import win32com.client
    shell = win32com.client.Dispatch("WScript.Shell")
    shell.SendKeys(easy_key_translation(key), 0)


@activity
def press_key_combination(first_key, second_key, third_key=None, force_pyautogui=False):
    """Press key combination

    Press a combination of two or three keys simultaneously. Make sure your keyboard is on US layout (standard QWERTY).

    Supported keys:
        ' ', '!', '"', '#', '$', '%', '&', "'", '(', ,')', '*', '+', ',', '-', '.', '/', '0', '1', '2', '3', '4', '5', '6', '7', '8', '9', ':', ';', '<','=', '>', '?', '@', '[', '\\', ']', '^', '_', '`', 'a', 'b', 'c', 'd', 'e','f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z', '{', '|', '}', '~', 'alt', 'backspace',  'ctrl', 'delete' 'downarrow', 'rightarrow', 'leftarrow', 'uparrow', 'enter', 'escape', 'f1', 'f2', f3', 'f4', 'f5', 'f6', 'f7', 'f8',  'f9', 'f10', 'f11', 'f12', 'f13', 'f14', 'f15', 'f16', 'home', 'insert', 'pagedown', 'pageup', 'help', 'printscreen', 'space', 'scrollock', 'tab', shift, 'win']

    :parameter first_key: First key to press
    :parameter second_key: Second key to press
    :parameter third_key: Third key to press, this is optional.
    :parameter force_pyautogui: Set parameter to true to force the use of pyautogui. This could help when certain keypresses do not work correctly.

    :return: Key combination

        :Example:

    >>> # Open notepad to illustrate typing
    >>> run('notepad.exe')
    >>> # Press 'ctrl + s' to prompt save window 
    >>> press_key_combination('ctrl', 's')

    Keywords
        keyboard, key combination, shortcut, typing, type, key, keystroke, hotkey, press, press key

    Icon
        las la-keyboard

    """

    import platform

    # Check if system is not running Windows
    if first_key == 'win' or second_key == 'win' or third_key == 'win':
        force_pyautogui = True
    if platform.system() != "Windows" or force_pyautogui:
        from pyautogui import hotkey

        if not third_key:
            return hotkey(first_key, second_key)
        if third_key:
            return hotkey(first_key, second_key, third_key)

    import win32com.client
    shell = win32com.client.Dispatch("WScript.Shell")
    key_combination = easy_key_translation(
        first_key) + easy_key_translation(second_key) + easy_key_translation(third_key)
    shell.SendKeys(easy_key_translation(key_combination), 0)


@activity
def type_text(text='', interval_seconds=0.01):
    """Type text

    Types text in the current active field by simulating keyboard typing.  Make sure your keyboard is on US layout (standard QWERTY).

    Supported keys: 
    ' ', '!', '"', '#', '$', '%', '&', "'", '(', ,')', '*', '+', ',', '-', '.', '/', '0', '1', '2', '3', '4', '5', '6', '7', '8', '9', ':', ';', '<','=', '>', '?', '@', '[', '\\', ']', '^', '_', '`', 'a', 'b', 'c', 'd', 'e','f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z', '{', '|', '}', '~', 'alt', 'backspace',  'ctrl', 'delete' 'downarrow', 'rightarrow', 'leftarrow', 'uparrow', 'enter', 'escape', 'f1', 'f2', f3', 'f4', 'f5', 'f6', 'f7', 'f8',  'f9', 'f10', 'f11', 'f12', 'f13', 'f14', 'f15', 'f16', 'home', 'insert', 'pagedown', 'pageup', 'help', 'printscreen', 'space', 'scrollock', 'tab', shift, 'win'

    :parameter text: Text in string format to type. Note that you can only press single character keys. Special keys like ":", "F1",... can not be part of the text argument.
    :parameter interval_seconds: Time in seconds between two keystrokes. Defautl value is 0.01 seconds.

    :return: Keystrokes

        :Example:

    >>> # Open notepad to illustrate typing
    >>> run('notepad.exe')
    >>> # Type a story
    >>> type_text('Why was the robot mad? \n They kept pushing his buttons!')

    Keywords
        keyboard, keystrokes, key combination, shortcut, typing, type, key, keystroke, hotkey, press, press key

    Icon
        las la-keyboard

    """

    import platform

    # Set keyboard layout for Windows platform
    if platform.system() != "Windows":
        from pyautogui import typewrite
        return typewrite(text, interval=interval_seconds)

    import win32com.client
    shell = win32com.client.Dispatch("WScript.Shell")
    import time
    for character in text:
        shell.SendKeys(easy_key_translation(character), 0)
        time.sleep(interval_seconds)


"""
Mouse
Icon: las la-mouse-pointer
"""


@activity
def get_mouse_position(delay=None, to_clipboard=False):
    """Get mouse coordinates

    Get the x and y pixel coordinates of current mouse position.
    These coordinates represent the absolute pixel position of the mouse on the computer screen. The x-coördinate starts on the left side and increases going right. The y-coördinate increases going down.

    :parameter delay: Delay in seconds before capturing mouse position.
    :parameter to_clipboard: Put the coordinates in the clipboard e.g. 'x=1, y=1'

    :return: Tuple with (x, y) coordinates

        :Example:

    >>> get_mouse_position()
    (314, 271)

    Keywords
        mouse, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-mouse
    """
    from pyautogui import position
    from time import sleep

    if delay:
        sleep(delay)

    coord = position()

    if to_clipboard:
        set_to_clipboard("x=" + str(coord[0]) + ", y=" + str(coord[1]))

    return coord[0], coord[1]


@activity
def display_mouse_position(duration=10):
    """Display mouse position

    Displays mouse position in an overlay. Refreshes every two seconds. Can be used to find mouse position of element on the screen. 
    These coordinates represent the absolute pixel position of the mouse on the computer screen. The x-coördinate starts on the left side and increases going right. The y-coördinate increases going down.

    :parameter duration: Duration to show overlay.

    :return: Overlay with (x, y) coordinates

        :Example:

    >>> display_mouse_position()

    Keywords
        mouse, osd, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        lars la-search-location
    """

    if duration < 1 or type(duration) != int:
        return

    from pyautogui import position
    from time import sleep

    duration_half = int(duration / 2)
    for i in range(0, duration_half, 2):
        coord = position()
        message = "x=" + str(coord[0]) + ", y=" + str(coord[1])
        display_osd_message(message, seconds=2)
        sleep(2)


@activity
def click(x=None, y=None):
    """Mouse click

    Clicks on a pixel position on the visible screen determined by x and y coordinates.

    :parameter x: X-coördinate
    :parameter y: Y-coördinate

    :return: Mouse click on (x, y) coordinates

        :Example:

    >>> click(x=100, y=100)

    Keywords
        mouse, osd, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-mouse-pointer
    """
    from pyautogui import click

    return click(x, y)


@activity
def double_click(x=None, y=None):
    """Double mouse click

    Double clicks on a pixel position on the visible screen determined by x and y coordinates.

    :parameter x: X-coördinate
    :parameter y: Y-coördinate

    :return: Double mouse click on (x, y) coordinates

        :Example:

    >>> double_click(x=100, y=100)

    Keywords
        mouse, osd, overlay, double, double click, doubleclick show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-mouse-pointer
    """
    from pyautogui import doubleClick

    return doubleClick(x, y)


@activity
def right_click(x=None, y=None):
    """Right click

    Right clicks on a pixel position on the visible screen determined by x and y coordinates.

    :parameter x: X-coördinate
    :parameter y: Y-coördinate

    :return: Right mouse click on (x, y) coordinates

        :Example:

    >>> right_click(x=100, y=100)

    Keywords
        mouse, osd, right click, right, rightclick, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-mouse-pointer
    """

    from pyautogui import rightClick

    return rightClick(x, y)


@activity
def move_mouse_to(x=None, y=None):
    """Move mouse

    Moves te pointer to a x-y position.

    :parameter x: X-coördinate
    :parameter y: Y-coördinate

    :return: Move mouse to (x, y) coordinates

        :Example:

    >>> move_mouse_to(x=100, y=100)

    Keywords
        mouse, osd, move mouse, right click, right, rightclick, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-arrows-alt
    """

    from pyautogui import moveTo

    return moveTo(x, y)


@activity
def move_mouse_relative(x=None, y=None):
    """Move mouse relative

    Moves the mouse an x- and y- distance relative to its current pixel position.

    :parameter x: X-coördinate
    :parameter y: Y-coördinate

    :return: Move mouse (x, y) coordinates

        :Example:

    >>> move_mouse_to(x=100, y=100)
    >>> wait(1)
    >>> move_mouse_relative(x=10, y=10)

    Keywords
        mouse, osd, move mouse, right click, right, rightclick, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-arrows-alt
    """

    from pyautogui import moveRel

    return moveRel(x, y)


@activity
def drag_mouse_to(x=None, y=None, button="left"):
    """Drag mouse

    Drag the mouse from its current position to a entered x-y position, while holding a specified button.

    :parameter x: X-coördinate
    :parameter y: Y-coördinate
    :parameter button: Button to hold while dragging. Options are 'left', 'middle' and 'right'. Standard value is 'left'.

    :return: Drag mouse (x, y) coordinates

        :Example:

    >>> move_mouse_to(x=100, y=100)
    >>> drag_mouse_to(x=1, y=1)

    Keywords
        mouse, osd, move mouse, right click, right, rightclick, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-arrows-alt
    """
    from pyautogui import dragTo

    return dragTo(x, y, 0.2, button=button)


"""
Image
Icon: las la-image
"""


@activity
def random_screen_snippet(size=100, path=None):
    """Random screen snippet

    Take a random square snippet from the current screen. Mainly for testing and/or development purposes.

    :parameter size: Size (width and height) in pixels for square snippet. Default value is 100 pixels
    :parameter path: Path where snippet will be saved. Default value is home directory with name 'random_screensnippet.jpg'

    :return: Path to snippet

        :Example:

    >>> random_screen_snippet()
    'C:\\Users\\<username>\\random_screensnippet.jpg'

    Keywords
        image, random, testing, screengrab, snippet

    Icon
        las la-crop-alt

    """
    import PIL.ImageGrab
    img = PIL.ImageGrab.grab()

    width, height = img.size

    import random
    random_left = random.randrange(1, width, 1)
    random_top = random.randrange(1, height, 1)

    left, top, right, bottom = random_left, random_top, random_left + size, random_top + size
    cropped = img.crop((left, top, right, bottom))

    if not path:
        import os
        path = os.path.join(os.path.expanduser(
            "~"), "random_screensnippet.jpg")
    cropped.save(path, "JPEG")

    return path


@activity
def take_screenshot(path=None):
    """Screenshot

    Take a screenshot of current screen.

    :parameter path: Path to save screenshot. Default value is home directory with name 'screenshot.jpg'.

    :return: Path to screenshot

        :Example:

    >>> take_screenshot
    'C:\\Users\\<username>\\screenshot.jpg'

    Keywords
        image, screenshot, printscreen,

    Icon
        las la-expand

    """
    import PIL.ImageGrab
    img = PIL.ImageGrab.grab()

    if not path:
        import os
        path = os.path.join(os.path.expanduser("~"), "screenshot.jpg")
    img.save(path, "JPEG")

    return path


@activity
def click_image(filename=None):
    """Click on image

    This function searches the screen for a match with template image and clicks directly in the middle. Note that this only finds exact matches.
    For a more advanced and robust vision detection method see Automagica AI functionality.

    :parameter filename: Path to the template image.

    :return: True if image was found and clicked on. False if template image was not found.

        :Example:

    >>> # Create a random snippet from current screen
    >>> # This is for illustration and can be replaced by template
    >>> snippet = random_screen_snippet(size=10)
    >>> # Click on the snippet
    >>> click_image(snippet)

    Keywords
        image matching, matching, vision, button, click, template, template matching.

    Icon
        las la-image

    """
    if not filename:
        return

    from pyautogui import locateCenterOnScreen, click

    try:
        x, y = locateCenterOnScreen(filename)
        click(x, y)
        return True
    except:
        return False


@activity
def double_click_image(filename=None):
    """Double click image

    Double click on similar image on the screen. This function searches the screen for a match with template image and doubleclicks directly in the middle.
    Note that this only finds exact matches. For a more advanced and robust vision detection method see Automagica AI functionality.

    :parameter filename: Path to the template image.

    :return: True if image was found and double clicked on. False if template image was not found.

        :Example: 

    >>> # Create a random snippet from current screen
    >>> # This is for illustration and can be replaced by template
    >>> snippet = random_screen_snippet(size=10)
    >>> # Click on the snippet
    >>> double_click_image(snippet)

    Keywords
        image matching, matching, vision, button, double click, template, template matching,click 

    Icon
        las la-image
    """
    if not filename:
        return

    from pyautogui import locateCenterOnScreen, click

    try:
        x, y = locateCenterOnScreen(filename)
        click(x, y, 2)
        return True
    except:
        return False


@activity
def right_click_image(filename=None):
    """Right click image

    Right click on similar image on the screen. This function searches the screen for a match with template image and right clicks directly in the middle.
    Note that this only finds exact matches. For a more advanced and robust vision detection method see Automagica AI functionality.

    :return: True if image was found and right clicked on. False if template image was not found.

        :Example:

    >>> # Create a random snippet from current screen
    >>> # This is for illustration and can be replaced by template
    >>> snippet = random_screen_snippet(size=10)
    >>> # Click on the snippet
    >>> right_click_image(snippet)

    Keywords
        image matching, matching, vision, button, right click, template, template matching, click

    Icon
        las la-image
    """

    if not filename:
        return

    from pyautogui import locateCenterOnScreen, rightClick

    try:
        x, y = locateCenterOnScreen(filename)
        rightClick(x, y, 2)
        return True
    except:
        return False


@activity
def locate_image_on_screen(filename=None):
    """Locate image on screen

    Find exact image on the screen. This function searches the screen for a match with template image and clicks directly in the middle.
    Note that this only finds exact matches. For a more advanced and robust vision detection method see Automagica AI functionality.

    :parameter filename: Path to the template image.

    :return: Tuple with (x, y) coordinates if image is found. None if image was not found

        :Example:

    >>> # Create a random snippet from current screen
    >>> # This is for illustration and can be replaced by template
    >>> snippet = random_screen_snippet(size=10)
    >>> # Click on the snippet
    >>> locate_image_on_screen(snippet)

    Keywords
        image matching, matching, vision, button, right click, template, template matching, click

    Icon
        las la-image
    """

    if not filename:
        return
    from pyautogui import locateCenterOnScreen

    try:
        x, y = locateCenterOnScreen(filename)
        return x, y
    except:
        return None


"""
Folder Operations
Icon: las la-folder-open
"""


@activity
def get_files_in_folder(path=None, extension=None, show_full_path=True, scan_subfolders=False):
    """List files in folder

    List all files in a folder (and subfolders)
    Checks all folders and subfolders for files. This could take some time for large repositories. 

    :parameter path: Path of the folder to retreive files from. Default folder is the home directory.
    :parameter extension: Optional filter on certain extensions, for example 'pptx', 'exe,' xlsx', 'txt', .. Default value is no filter.
    :parameter show_full_path: Set this to True to show full path, False will only show file or dirname. Default is True
    :scan_subfolders: Boolean to scan subfolders or not. Not that depending on the folder and hardware this activity could take some time if scan_subfolders is set to True

    :return: List of files with their full path

        :Example:

    >>> # List all files in the homedirectory
    >>> get_files_in_folder()
    ['C:\\Users\\<username>\\file1.jpg', 'C:\\Users\\<username>\\file2.txt', ... ]

    Keywords
        folder, files, explorer, nautilus, folder, file, create folder, get files, list files, all files, overview, get files

    Icon 
        las la-search
    """
    import os

    if not path:
        path = os.path.expanduser("~")

    if scan_subfolders:
        paths = []
        for dirpath, _, filenames in os.walk(path):
            for f in filenames:
                full_path = os.path.abspath(os.path.join(dirpath, f))
                if extension:
                    if not full_path.endswith(extension):
                        continue
                if show_full_path:
                    paths.append(full_path)
                else:
                    paths.append(f)
        return paths

    if not scan_subfolders:
        paths = []
        for item in os.listdir(path):
            full_path = os.path.abspath(os.path.join(path, item))
            if extension:
                if not full_path.endswith(extension):
                    continue
            if show_full_path:
                paths.append(full_path)
            else:
                paths.append(item)
        return paths


@activity
def create_folder(path=None):
    """Create folder

    Creates new folder at the given path.

    :parameter path: Full path of folder that will be created. If no path is specified a folder called 'new_folder' will be made in home directory. If this folder already exists 8 random characters will be added to the name.

    :return: Path to new folder as string

        :Example:

    >>> # Create folder in the home directory
    >>> create_folder()
    'C:\\Users\\<username>\\new_folder'

    Keywords
        create folder, folder, folders, make folder, new folder, folder manipulation, explorer, nautilus

    Icon
        las la-folder-plus
    """
    import os

    if not path:
        path = os.path.join(os.path.expanduser("~"), "new_folder")

    if not os.path.exists(path):
        os.makedirs(path)
        return path

    from uuid import uuid4
    folder_name = os.path.basename(path)
    extended_folder_name = folder_name + '_' + str(uuid4())[:4]
    parent_dir = os.path.abspath(os.path.join(path, os.pardir))
    path = os.path.abspath(os.path.join(parent_dir, extended_folder_name))
    os.makedirs(path)

    return path


@activity
def rename_folder(input_path, new_name=None):
    """Rename folder

    Rename a folder

    :parameter path: Full path of folder that will be renamed
    :parameter new_name: New name of the folder e.g. 'new_folder'. By default folder will be renamed to original folder name with '_renamed' added to the folder name.

    :return: Path to renamed folder as a string. None if folder could not be renamed.

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Rename the folder
    >>> rename_folder(testfolder, new_name='testfolder_brand_new_name')
    'C:\\Users\\<username>\\testfolder_brand_new_name'

    Keywords
        folder, rename, rename folder, organise folder, folders, folder manipulation, explorer, nautilus

    Icon
        las la-folder

    """
    import os

    if not os.path.exists(input_path):
        return None

    if not new_name:
        folder_name = os.path.basename(input_path)
        new_name = folder_name + '_renamed'

    parent_dir = os.path.abspath(os.path.join(input_path, os.pardir))
    renamed_dir = os.path.join(parent_dir, new_name)

    if os.path.exists(renamed_dir):
        return None

    os.rename(input_path, renamed_dir)

    return renamed_dir


@activity
def show_folder(path=None):
    """Open a folder

    Open a folder with the default explorer.

    :parameter path: Full path of folder that will be opened. Default value is the home directory

    :return: Path to opend folder as a string

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Open folder
    >>> show_folder(testfolder)
    'C:\\Users\\<username>\\new_folder'


    Keywords
        folder, open, open folder, explorer, nautilus

    Icon
        las la-folder-open

    """
    import os

    if not path:
        path = os.path.expanduser("~")

    if os.path.isdir(path):
        os.startfile(path)

    return path


@activity
def move_folder(from_path, to_path):
    """Move a folder

    Moves a folder from one place to another.
    If the new location already contains a folder with the same name, a random 4 character uid is added to the name.

    :parameter fom_path: Full path to the source location of the folder
    :parameter new_path: Full path to the destination location of the folder.

     :return: Path to new location of folder as a string. None if folder could not be moved.

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> # If no new_folder exists in home dir this will be called new_folder
    >>> testfolder = create_folder()
    >>> # Make a second new folder
    >>> # Since new_folder already exists this folder will get a random id added (in this case abc1)
    >>> testfolder_2 = create_folder()
    >>> # Move testfolder in testfolder_2
    >>> move_folder(testfolder, testfolder_2)
    'C:\\Users\\<username>\\new_folder_abc1\\new_folder'

    Keywords
        folder, move, move folder, explorer, nautilus, folder manipulation

    Icon
        las la-folder
    """
    from uuid import uuid4
    import os
    import shutil

    from_path_folder = os.path.basename(from_path)
    if os.path.isdir(from_path):
        if not os.path.isdir(to_path):
            shutil.move(from_path, to_path)
            return os.path.join(to_path, from_path_folder)
        elif os.path.isdir(to_path):
            to_path_folder_name = os.path.basename(to_path)
            to_path_folder_name = to_path_folder_name + str(uuid4())[:4]
            to_path_base_name = os.path.abspath(
                os.path.join(to_path, os.pardir))
            to_path = os.path.join(to_path_base_name, to_path_folder_name)
            shutil.move(from_path, to_path)
            return os.path.join(to_path, from_path_folder)
    return None


@activity
def remove_folder(path, allow_root=False, delete_read_only=True):
    """Remove folder

    Remove a folder including all subfolders and files. For the function to work optimal, all files and subfolders in the main targetfolder should be closed.

    :parameter path: Full path to the folder that will be deleted
    :parameter allow_root: Allow paths with an arbitrary length of 10 characters or shorter to be deleted. Default value is False.

    :return: Path to deleted folder as a string

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Check if folder exists
    >>> print( folder_exists(testfolder) ) # Should print True
    >>> # Remove folder
    >>> remove_folder(testfolder)
    >>> # Check again if folder exists
    >>> folder_exists(testfolder)
    False

    Keywords
        folder, delete folder, delete, nautilus, folder manipulation, explorer, delete folder, remove, remove folder

    Icon
        las la-folder-minus

    """
    import os
    import shutil

    if len(path) > 10 or allow_root:
        if os.path.isdir(path):
            shutil.rmtree(path, ignore_errors=delete_read_only)
            return path


@activity
def empty_folder(path, allow_root=False):
    """Empty folder
    
    Remove all contents from a folder
    For the function to work optimal, all files and subfolders in the main targetfolder should be closed.

    :parameter path: Full path to the folder that will be emptied
    :parameter allow_root: Allow paths with an arbitrary length of 10 characters or shorter to be emptied. Default value is False.

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Make new textfile in this folder
    >>> import os
    >>> textfile_location = os.path.join(testfolder, 'testfile.txt')
    >>> make_textfile(output_path=textfile_location )
    >>> # Print all files in the testfolder
    >>> print( get_files_in_folder(testfolder) ) # Should show 
    >>> # Empty the folder
    >>> empty_folder(testfolder)
    >>> # Check what is in the folder
    >>> get_files_in_folder(testfolder)
    []

    Keywords
        folder, empty folder, delete, empty, clean, clean folder, nautilus, folder manipulation, explorer, delete folder, remove, remove folder

    Icon
        las la-folder-minus
    """
    import os

    if len(path) > 10 or allow_root:
        if os.path.isdir(path):
            for root, dirs, files in os.walk(path, topdown=False):
                for name in files:
                    os.remove(os.path.join(root, name))
                for name in dirs:
                    os.rmdir(os.path.join(root, name))


@activity
def folder_exists(path):
    """Checks if folder exists

    Check whether folder exists or not, regardless if folder is empty or not.

    :parameter path: Full path to folder

    :return: Boolean

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Check if folder exists
    >>> folder_exists(testfolder)
    True

    Keywords
        folder, folder exists, nautilus, explorer, folder manipulation, files

    Icon
        las la-folder
    """
    import os

    return os.path.isdir(path)


@activity
def copy_folder(from_path, to_path=None):
    """Copy a folder

    Copies a folder from one place to another.
    If the new location already contains a folder with the same name, a random 4 character id is added to the name.

    :parameter old_path: Full path to the source location of the folder
    :parameter new_path: Full path to the destination location of the folder. If no path is specified folder will get copied in the from_path directory

    :return: Path to new folder as string

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Copy this folder
    >>> # Since new_folder already exists in home dir this folder will get a random id added (in this case abc1)
    >>> copy_folder(testfolder)

    Keywords
        folder, move, move folder, explorer, nautilus, folder manipulation

    Icon
        lar la-folder
    """
    from uuid import uuid4
    import os
    import shutil

    if not to_path:
        to_path = from_path

    from_path_folder = os.path.basename(from_path)
    if os.path.isdir(from_path):
        if not os.path.isdir(to_path):
            shutil.copytree(from_path, to_path)
            return os.path.join(to_path, from_path_folder)
        elif os.path.isdir(to_path):
            to_path_folder_name = os.path.basename(
                to_path) + '_copy_' + str(uuid4())[:4]
            to_path_base_name = os.path.abspath(
                os.path.join(to_path, os.pardir))
            to_path = os.path.join(to_path_base_name, to_path_folder_name)
            shutil.copytree(from_path, to_path)
            return os.path.join(to_path, from_path_folder)
    return None


@activity
def zip_folder(path, new_path=None):
    """Zip

     Zia folder and it's contents. Creates a .zip file. 

    :parameter path: Full path to the source location of the folder that will be zipped
    :parameter new_path: Full path to save the zipped folder. If no path is specified a folder with the original folder name plus 4 random characters

    :return: Path to zipped folder

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Zip this folder
    >>> zip_folder(testfolder)

    Keywords
        zip, zipping, winrar, rar, 7zip, compress, unzip

    Icon
        las la-archive
    """
    import zipfile
    import os
    import shutil
    from uuid import uuid4

    if not new_path:
        from uuid import uuid4
        new_path = path + "_" + str(uuid4())[:4]
    if os.path.isdir(path):
        shutil.make_archive(new_path, "zip", path)
    return new_path


@activity
def unzip(path, to_path=None):
    """Unzip

    Unzips a file or folder from a .zip file.

    :parameter path: Full path to the source location of the file or folder that will be unzipped
    :parameter to_path: Full path to save unzipped contents. If no path is specified the unzipped contents will be stored in the same directory as the zipped file is located. 

    :return: Path to unzipped folder

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Zip this folder
    >>> zipped_folder = zip_folder(testfolder)
    >>> # Unzip this folder
    >>> unzip(zipped_folder)

    Keywords
        zip, zipping, winrar, rar, 7zip, compress, unzip

    Icon
        las la-archive
    """
    import zipfile
    import os
    import shutil

    if os.path.exists(path):
        zipp = zipfile.ZipFile(path)
        if not to_path:
            to_path = from_path_folder = os.path.basename(path)
            zipp.extractall(to_path)
        elif os.path.isdir(to_path):
            zipp.extractall(to_path)
        zipp.close()
    return to_path


"""
Delay
Icon: las la-hourglass
"""


@activity
def wait(seconds=1):
    """Wait

    Make the robot wait for a specified number of seconds. Note that this activity is blocking. This means that subsequent activities will not occur until the the specified waiting time has expired.

    :parameter seconds: Time in seconds to wait

        :Example:

    >>> print('Start the wait')
    >>> wait()
    >>> print('The wait is over')

    Keywords
        wait, sleep, time, timeout, time-out, hold, pause

    Icon
        las la-hourglass        
    """
    from time import sleep

    sleep(seconds)


@activity
def wait_for_image(path=None, timeout=60):
    """Wait for image

    Waits for an image to appear on the screen
    Note that this activity waits for an exact match of the template image to appear on the screen. 
    Small variations, such as color or resolution could result in a mismatch.

    :parameter path: Full or relative path to the template image.
    :parameter timeout: Maximum time in seconds to wait before continuing. Default value is 60 seconds.

        :Example:

    >>> # Create a random snippet from current screen
    >>> # This is for illustration and can be replaced by template
    >>> snippet = random_screen_snippet(size=10)
    >>> # Wait for the snippet to be visible
    >>> wait_for_image(snippet)

    Keywords
        image matching, wait, pause, vision, template, template matching

    Icon
        las la-hourglass
    """

    from pyautogui import locateCenterOnScreen
    from time import sleep

    for _ in range(timeout):
        try:
            locateCenterOnScreen(path)
            break
        except TypeError:
            sleep(1)


@activity
def wait_folder_exists(path, timeout=60):
    """Wait for folder

    Waits until a folder exists.
    Not that this activity is blocking and will keep the system waiting.

    :parameter path: Full path to folder.
    :parameter timeout: Maximum time in seconds to wait before continuing. Default value is 60 seconds.

        :Example:

    >>> # Create a random folder
    >>> testfolder = create_folder()
    >>> # Wait for the snippet to be visible
    >>> wait_folder_exists(testfolder)

    Keywords
        image matching, wait, pause, vision, template, template matching

    Icon
        las la-hourglass
    """
    from time import sleep
    import os

    while not os.path.exists(path):
        sleep(1)
    return

    for _ in range(timeout):
        if os.path.exists(path):
            break
            sleep(1)


"""
Word Application
Icon: las la-file-word
"""


class Word:
    @activity
    def __init__(self, visible=True, file_path=None):
        """Start Word Application

        For this activity to work, Microsoft Office Word needs to be installed on the system.

        :parameter visible: Show Word in the foreground if True or hide if False, defaults to True.
        :parameter path: Enter a path to open Word with an existing Word file. If no path is specified a document will be initialized, this is the default value.

            :Example:

        >>> word = Word()

        Keywords
            word, editor, text, text edit, office, document, microsoft word, doc, docx

        Icon
            lar la-file-word

        """
        self.file_path = file_path

        self.app = self._launch()
        self.app.Visible = visible

    def _launch(self):
        """Utility function to create the Word application scope object

        :return: Application object (win32com.client)
        """
        try:
            import win32com.client

            app = win32com.client.gencache.EnsureDispatch("Word.Application")
            #app = win32com.client.dynamic.Dispatch("Word.Application")

        except:
            raise Exception(
                "Could not launch Word, do you have Microsoft Office installed on Windows?"
            )

        if self.file_path:
            app.Documents.Open(self.file_path)
        else:
            app.Documents.Add()

        return app

    @activity
    def append_text(self, text):
        """Append text

        Append text at end of Word document.

        :parameter text: Text to append to document

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')

        Keywords
            word, editor, text, text edit, office, document, microsoft word, doc, docx

        Icon
            lar la-file-word
        """
        import win32com.client

        wc = win32com.client.constants
        self.app.Selection.EndKey(Unit=wc.wdStory)
        self.app.Selection.TypeText(text)

    @activity
    def replace_text(self, placeholder_text, replacement_text):
        """Replace text

        Can be used for example to replace arbitrary placeholder value. For example when 
        using template document, using 'XXXX' as a placeholder. Take note that all strings are case sensitive.

        :parameter placeholder_text: Placeholder text value (string) in the document, this will be replaced, e.g. 'Company Name'
        :parameter replacement_text: Text (string) to replace the placeholder values with. It is recommended to make this unique to avoid wrongful replacement, e.g. 'XXXX_placeholder_XXX'

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')
        >>> word.replace_text('sample', 'real')

        Keywords
            word, replace, text, template

        Icon
            lar la-file-word
        """

        self.app.Selection.GoTo(0)
        self.app.Selection.Find.Text = placeholder_text
        self.app.Selection.Find.Replacement.Text = replacement_text
        self.app.Selection.Find.Execute(Replace=2, Forward=True)

    @activity
    def read_all_text(self, return_as_list=False):
        """Read all text

        Read all the text from a document

        :parameter return_as_list: Set this paramater to True to return text as a list of strings. Default value is False.

        :return: Text from the document

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')
        >>> word.replace_text('sample', 'real')
        >>> word.read_all_text()
        'This is real text'

        Keywords
            word, extract, text, document

        Icon
            lar la-file-word
        """

        if return_as_list:
            return self.app.ActiveDocument.Content.Text.split('\r')
        return self.app.ActiveDocument.Content.Text

    @activity
    def export_to_pdf(self, path=None):
        """Export to PDF

        Export the document to PDF

        :parameter path: Output path where PDF file will be exported to. Default path is home directory with filename 'pdf_export.pdf'.

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')
        >>> word.replace_text('sample', 'real')
        >>> word.export_to_pdf('output.pdf')

        Keywords
            word, pdf, document, export, save as

        Icon
            lar la-file-pdf

        """

        if not path:
            import os
            path = os.path.expanduser("~") + '/pdf_export.pdf'

        self.app.ActiveDocument.ExportAsFixedFormat(OutputFileName=path,
                                                    ExportFormat=17,
                                                    OpenAfterExport=False,
                                                    OptimizeFor=0,
                                                    CreateBookmarks=1,
                                                    DocStructureTags=True
                                                    )

    @activity
    def export_to_html(self, path=None):
        """Export to HTML

        Export to HTML

        :parameter path: Output path where HTML file will be exported to. Default path is home directory with filename 'html_export.html'.

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')
        >>> word.replace_text('sample', 'real')
        >>> word.export_to_html('output.html')

        Keywords
            word, html, document, export, save as

        Icon
            las la-html5

        """
        if not path:
            import os
            path = os.path.expanduser("~") + '/pdf_export.pdf'

        import win32com.client

        wc = win32com.client.constants
        word.app.ActiveDocument.WebOptions.RelyOnCSS = 1
        word.app.ActiveDocument.WebOptions.OptimizeForBrowser = 1
        word.app.ActiveDocument.WebOptions.BrowserLevel = 0
        word.app.ActiveDocument.WebOptions.OrganizeInFolder = 0
        word.app.ActiveDocument.WebOptions.UseLongFileNames = 1
        word.app.ActiveDocument.WebOptions.RelyOnVML = 0
        word.app.ActiveDocument.WebOptions.AllowPNG = 1
        word.app.ActiveDocument.SaveAs(
            FileName=path, FileFormat=wc.wdFormatHTML)

    @activity
    def set_footers(self, text):
        """Set footers
        
        Set the footers of the document

        :parameter text: Text to put in the footer

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.set_footers('This is a footer!')


        Keywords
            word, footer, footers

        Icon
            las la-heading
        """
        for section in self.app.ActiveDocument.Sections:
            for footer in section.Footers:
                footer.Range.Text = text

    @activity
    def set_headers(self, text):
        """Set headers

        Set the headers of the document

        :parameter text: Text to put in the header

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.set_headers('This is a header!')

        Keywords
            word, header, headers

        Icon
            las la-subscript
        """
        for section in self.app.ActiveDocument.Sections:
            for footer in section.Headers:
                footer.Range.Text = text


"""
Word File
Icon: las la-file-word
"""

class WordFile:
    @activity
    def __init__(self, file_path=None):
        """Read and Write Word files

        These activities can read, write and edit Word (docx) files without the need of having Word installed. 
        Note that, in contrary to working with the :func: 'Word' activities, a file get saved directly after manipulation.

        :parameter file_path: Enter a path to open Word with an existing Word file. If no path is specified a 'document.docx' will be initialized in the home directory, this is the default value. If a document with the same name already exists the file will be overwritten.

            :Example:

        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.read_all_text()
        'Some sample text'

        Keywords
            word, read, text, file

        Icon
            las la-file-word

        """

        self.file_path = file_path

        self.app = self._launch()

    def _launch(self):
        from docx import Document
        import os

        if self.file_path:
            if not os.path.exists(self.file_path):
                document = Document(self.file_path)
        else:
            path = os.path.expanduser("~") + '\document.docx'
            document = Document()
            document.save(path)
            self.file_path = path

    @activity
    def read_all_text(self, return_as_list=False):
        """Read all text

        Read all the text from the document

        :parameter return_as_list: Set this paramater to True to return text as a list of strings. Default value is False.

        :return: Text of the document

            :Example:

        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.read_all_text()
        'Some sample text'

        Keywords
            word, read, text, file

        Icon
            las la-file-word
        """

        from docx import Document

        document = Document(self.file_path)
        text = []
        for paragraph in document.paragraphs:
            text.append(paragraph.text)

        if return_as_list:
            return text
        return '\r'.join(map(str, text))

    @activity
    def append_text(self, text, auto_save=True):
        """Append text

        Append text at the end of the document

        :parameter text: Text to append
        :parameter auto_save: Save document after performing activity. Default value is True

            :Example:

        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')

        Keywords
            word, append text, add text

        Icon
            las la-file-word
        """
        from docx import Document

        document = Document(self.file_path)
        document.add_paragraph(text)

        if auto_save:
            document.save(self.file_path)

    @activity
    def save(self):
        """Save

        Save document

            :Example:

        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.save()

        Keywords
            word, save, store

        Icon
            las la-file-word
        """
        document.save(self.file_path)

    @activity
    def save_as(self, path):
        """Save as

            :Example:
            
        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.save_as('document.docx')

        Keywords
            word, save as, store

        Icon
            las la-file-word
        """
        document.save(path)

    @activity
    def set_headers(self, text, auto_save=True):
        """Set headers

        Set headers of Word document

        :parameter text: Text to put in the header
        :parameter auto_save: Save document after performing activity. Default value is True

            :Example:
            
        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.set_headers('This is a header')

        Keywords
            word, header text

        Icon
            las la-file-word
        """
        from docx import Document

        document = Document(self.file_path)
        document.add_heading(text)

        if auto_save:
            document.save(self.file_path)

    @activity
    def replace_text(self, placeholder_text, replacement_text, auto_save=True):
        """Replace all

        Replaces all occurences of a placeholder text in the document with a replacement text.

        Can be used for example to replace arbitrary placeholder value. 
        For example when using template slidedeck, using 'XXXX' as a placeholder.
        Take note that all strings are case sensitive.

        :parameter placeholder_text: Placeholder text value (string) in the document, this will be replaced, e.g. 'Company Name'
        :parameter replacement_text: Text (string) to replace the placeholder values with. It is recommended to make this unique to avoid wrongful replacement, e.g. 'XXXX_placeholder_XXX'
        :parameter auto_save: Save document after performing activity. Default value is True

            :Example:
            
        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.replace_text('sample', 'real')

        Keywords
            word, replace text, template

        Icon
            las la-file-word
        """
        from docx import Document

        document = Document(self.file_path)
        for paragraph in document.paragraphs:
            paragraph.text = paragraph.text.replace(
                placeholder_text, replacement_text)

        if auto_save:
            document.save(self.file_path)


"""
Outlook Application
Icon: las la-envelope
"""


class Outlook:
    @activity
    def __init__(self, account_name=None):
        """Start Outlook Application

        For this activity to work, Outlook needs to be installed on the system.

            :Example:
            
        >>> outlook = Outlook()

        Keywords
            outlook, send e-mail, send mail

        Icon
            las la-mail-bulk
        
        """
        self.app = self._launch()
        self.account_name = account_name

    def _launch(self):
        """Utility function to create the Outlook application scope object

        :return: Application object (win32com.client)
        """
        try:
            import win32com.client

            app = win32com.client.gencache.EnsureDispatch(
                "outlook.application")

        except:
            raise Exception(
                "Could not launch Outlook, do you have Microsoft Office installed on Windows?"
            )

        return app

    @activity
    def send_mail(self, to_address, subject="", body="", html_body=None, attachment_paths=None):
        """Send e-mail

        Send an e-mail using Outlook

        :parameter to_address: The e-mail address the e-mail should be sent to
        :parameter subject: The subject of the e-mail
        :parameter body: The text body contents of the e-mail
        :parameter html_body: The HTML body contents of the e-mail (optional)
        :parameter attachment_paths: List of file paths to attachments

            :Example:
            
        >>> outlook = Outlook()
        >>> outlook.send_mail('test@test.com', subject='Hello world', body='Hi there')

        Keywords
            outlook, send e-mail, send mail

        Icon
            las la-mail-bulk
        """
        # mapi = self.app.GetNamespace("MAPI")

        # Create a new e-mail
        mail = self.app.CreateItem(0)

        mail.To = to_address
        mail.Subject = subject
        mail.Body = body

        if html_body:
            mail.HTMLBody = html_body

        # Add attachments
        if attachment_paths:
            for attachment_path in attachment_paths:
                mail.Attachments.Add(attachment_path)

        # Send the e-mail
        mail.Send()

    @activity
    def get_folders(self, limit=999):
        """Retrieve folders

        Retrieve list of folders from Outlook

        :parameter limit: Maximum number of folders to retrieve

            :Example:
            
        >>> outlook = Outlook()
        >>> outlook.get_folders()
        ['Inbox', 'Sent', ...]

        Keywords
            outlook, get folders, list folders

        Icon
            las la-mail-bulk
        """

        folders = []

        if self.account_name:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(self.account_name).Folders
        else:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(1).Folders
        for folder in found_folders:
            name = folder.Name
            folders.append(name)

        return folders

    @activity
    def get_mails(self, folder_name="Inbox", fields=None):
        """Retrieve e-mails

        Retrieve list of messages from Outlook

        :parameter folder_name: Name of the Outlook folder, can be found using `get_folders`.
        :parameter limit: Number of messages to retrieve
        :parameter fields: Fields (properties) of e-mail messages to give, requires tupl Stadard is 'Subject', 'Body', 'SentOn' and 'SenderEmailAddress'.

        :return: List of dictionaries containing the e-mail messages with from, to, subject, body and html.

            :Example:
            
        >>> outlook = Outlook()
        >>> outlook.get_mails()
        [
            {
                'Subject': 'Hello World!',
                'Body' : 'This is an e-mail',
                'SenderEmailAddress': 'from@test.com'
            }
        ]

        Keywords
            outlook, retrieve e-mail, receive e-mails, process e-mails, get mails

        Icon
            las la-mail-bulk
        """

        if not fields:
            fields = ("Subject", "Body", "SenderEmailAddress")

        messages = []

        if self.account_name:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(self.account_name).Folders
        else:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(1).Folders
        for folder in found_folders:
            name = folder.Name
            if name == folder_name:
                break
        else:
            raise Exception(
                "Could not find the folder with name '{}'.".format(folder_name)
            )

        # Loop over the items in the folder
        for item in folder.Items:
            message = {}

            for key in fields:
                try:
                    message[key] = getattr(item, key)
                except AttributeError:
                    pass

            messages.append(message)

        return messages

    @activity
    def delete_mails(self,folder_name="Inbox",limit=0,subject_contains="",body_contains="",sender_contains=""):
        """Delete e-mails

        Deletes e-mail messages in a certain folder. Can be specified by searching on subject, body or sender e-mail.

        :parameter folder_name: Name of the Outlook folder, can be found using `get_folders`
        :parameter limit: Maximum number of e-mails to delete in one go
        :parameter subject_contains: Only delete e-mail if subject contains this
        :parameter body_contains: Only delete e-mail if body contains this
        :parameter sender_contains: Only delete e-mail if sender contains this

            :Example:

        >>> outlook = Outlook()
        >>> outlook.delete_mails(subject_contains='hello')

        Keywords
            outlook, remove e-mails, delete mail, remove mail

        Icon
            las la-mail-bulk

        """
        # Find the appropriate folder
        if self.account_name:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(self.account_name).Folders
        else:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(1).Folders
        for folder in found_folders:
            name = folder.Name
            if name == folder_name:
                break
        else:
            raise Exception(
                "Could not find the folder with name '{}'.".format(folder_name)
            )

        # Loop over the items in the folder
        for i, item in enumerate(folder.Items):

            if limit:
                if i > limit:
                    break

            if subject_contains in item.Subject:
                if body_contains in item.Body:
                    if sender_contains in item.SenderEmailAddress:
                        item.Delete()

    @activity
    def move_mails(self,source_folder_name="Inbox",target_folder_name="Archive",limit=0,subject_contains="",body_contains="",sender_contains=""):
        """Move e-mails

        Move e-mail messages in a certain folder. Can be specified by searching on subject, body or sender e-mail.

        :parameter source_folder_name: Name of the Outlook source folder from where e-mails will be moved, can be found using `get_folders`
        :parameter target_folder_name: Name of the Outlook destination folder to where e-mails will be moved, can be found using `get_folders`
        :parameter limit: Maximum number of e-mails to move in one go
        :parameter subject_contains: Only move e-mail if subject contains this
        :parameter body_contains: Only move e-mail if body contains this
        :parameter sender_contains: Only move e-mail if sender contains this

            :Example:

        >>> outlook = Outlook()
        >>> outlook.move_mails(subject_contains='move me')

        Keywords
            outlook, move e-mail, move e-mail to folder

        Icon
            las la-mail-bulk
        """
        # Find the appropriate source folder
        if self.account_name:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(self.account_name).Folders
        else:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(1).Folders
        for source_folder in found_folders:
            name = source_folder.Name
            if name == source_folder_name:
                break
        else:
            raise Exception(
                "Could not find the folder with name '{}'.".format(
                    source_folder_name)
            )

        # Find the appropriate target folder
        if self.account_name:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(self.account_name).Folders
        else:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(1).Folders
        for target_folder in found_folders:
            name = target_folder.Name
            if name == target_folder_name:
                break
        else:
            raise Exception(
                "Could not find the folder with name '{}'.".format(
                    target_folder_name)
            )

        # Loop over the items in the folder
        for i, item in enumerate(source_folder.Items):

            if limit:
                if i > limit:
                    break

            if subject_contains in item.Subject:
                if body_contains in item.Body:
                    if sender_contains in item.SenderEmailAddress:
                        item.Move(target_folder)

    @activity
    def save_attachments(self, folder_name="Inbox", target_folder_path=None):
        """Save attachments

        :parameter folder_name: Name of the Outlook folder, can be found using `get_folders`.
        :parameter target_folder_path: Path where attachments will be saved. Default is the home directory.

        :return: List of paths to saved attachments.

            :Example:
            
        >>> outlook = Outlook()
        >>> outlook.save_attachments()
        ['Attachment.pdf', 'Signature_image.jpeg']

        Keywords
            outlook, save attachments, download attachments, extract attachments

        Icon
            las la-mail-bulk
        """
        import os

        paths = []

        # Set to user home if no path specified
        if not target_folder_path:
            target_folder_path = os.path.expanduser("~")

        # Find the appropriate folder
        if self.account_name:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(self.account_name).Folders
        else:
            found_folders = self.app.GetNamespace(
                "MAPI").Folders.Item(1).Folders
        for folder in found_folders:
            name = folder.Name
            if name == folder_name:
                break
        else:
            raise Exception(
                "Could not find the folder with name '{}'.".format(folder_name)
            )
        # Loop over the items in the folder
        for item in folder.Items:
            for attachment in item.Attachments:
                path = os.path.join(target_folder_path, attachment.FileName)
                attachment.SaveAsFile(path)
                paths.append(path)

        return paths

    @activity
    def get_contacts(self, fields=None):
        """Retrieve contacts

        :parameter fields: Fields can be specified as a tuple with their exact names. Standard value is None returning "LastName", "FirstName" and "Email1Address".

        :return: List of dictionaries containing the contact details.

            :Example:

        >>> outlook = Outlook()
        >>> outlook.get_contacts()
        [
            {
                'LastName': 'Doe',
                'FirstName' : 'John',
                'Email1Address': 'john@test.com'
            }
        ]

        Keywords
            outlook, get contacts, download contacts, rolodex

        Icon
            las la-mail-bulk
        """
        import win32com.client

        if not fields:
            fields = ("LastName", "FirstName", "Email1Address")

        contacts = []

        mapi = self.app.GetNamespace("MAPI")

        data = mapi.GetDefaultFolder(
            win32com.client.constants.olFolderContacts)

        for item in data.Items:
            if item.Class == win32com.client.constants.olContact:
                contact = {}
                for key in item._prop_map_get_:
                    if key in fields:
                        if isinstance(getattr(item, key), (int, str)):
                            contact[key] = getattr(item, key)
                contacts.append(contact)

        return contacts

    @activity
    def add_contact(self, email, first_name="", last_name=""):
        """Add a contact

        Add a contact to Outlook contacts

        :parameter email: The e-mail address for the contact
        :parameter first_name: First name for the contact (optional)
        :parameter last_name: Last name for the contact (optional)

            :Example:

        >>> outlook = Outlook()
        >>> outlook.add_contact('sales@automagica.com')

        Keywords
            outlook, create contact, add contact

        Icon
            las la-mail-bulk
        """

        # Create a new contact
        contact = self.app.CreateItem(2)

        contact.Email1Address = email

        if first_name:
            contact.FirstName = first_name

        if last_name:
            contact.LastName = last_name

        contact.Save()

    @activity
    def quit(self):
        """Quit

        Close the Outlook application

            :Example:
            
        >>> outlook = Outlook()
        >>> outlook.quit()

        Keywords
            outlook, close, quit

        Icon
            las la-mail-bulk
        """
        self.app.Application.Quit()


"""
Excel Application
Icon: las la-file-excel
"""


class Excel:
    @activity
    def __init__(self, visible=True, file_path=None):
        """Start Excel Application

        For this activity to work, Microsoft Office Excel needs to be installed on the system.

        :parameter visible: Show Excel in the foreground if True or hide if False, defaults to True.
        :parameter path: Enter a path to open Excel with an existing Excel file. If no path is specified a workbook will be initialized, this is the default value.

            :Example:

        >>> # Open Excel
        >>> excel = Excel()

        Keywords
            excel, add worksheet, add tab

        Icon
            las la-file-excel

        """
        self.file_path = file_path

        self.app = self._launch()
        self.app.Visible = visible

    def _launch(self):
        """Utility function to create the Excel application scope object

        :return: Application object (win32com.client)
        """
        try:
            import win32com.client

            app = win32com.client.gencache.EnsureDispatch("Excel.Application")

        except:
            raise Exception(
                "Could not launch Excel, do you have Microsoft Office installed on Windows?")

        if self.file_path:
            app.Workbooks.Open(self.file_path)
        else:
            app.Workbooks.Add()

        self.workbook = app.ActiveWorkbook

        return app

    @activity
    def add_worksheet(self, name=None):
        """Add worksheet

        Adds a worksheet to the current workbook

        :parameter workbook: Workbook object which is retrieved with either new_workbook or open_workbook
        :parmeter name: Give the sheet a name (optional)

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add a worksheet
        >>> excel.add_worksheet('My Example Worksheet')

        Keywords
            excel, add worksheet, add tab, insert worksheet, new worksheet

        Icon
            las la-file-excel
        
        """
        worksheet = self.workbook.Worksheets.Add()
        if name:
            worksheet.Name = name

    @activity
    def activate_worksheet(self, name):
        """Activate worksheet

        Activate a worksheet in the current Excel document by name

        :parameter name: Name of the worksheet to activate

            :Example:

        >>> # Open Excel   
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Add another worksheet
        >>> excel.add_worksheet('Another Worksheet')
        >>> # Activate the first worksheet
        >>> excel.activate_worksheet('My Example Worksheet)
        

        Keywords
            excel, activate worksheet, set worksheet, select worksheet, select tab, activate tab

        Icon
            las la-file-excel
        
        """
        for worksheet in self.workbook.Worksheets:
            if worksheet.Name == name:
                worksheet.Activate()

    @activity
    def save(self):
        """Save
        
        Save the current workbook

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Save the workbook to My Documents
        >>> excel.save()

        Keywords
            excel, save, store

        Icon
            las la-file-excel
        """
        self.workbook.Save()

    @activity
    def save_as(self, path):
        """Save as

        Save the current workbook to a specific path

        :parameter path: Path where workbook will be saved. Default is home directory and filename 'workbook.xlsx'

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Save the workbook to the current working directory
        >>> excel.save_as('output.xlsx')

        Keywords
            excel, save as, export

        Icon
            las la-file-excel
        
        """
        if not path:
            import os
            path = os.path.expanduser("~") + '\workbook.xlsx'

        self.app.DisplayAlerts = False
        self.workbook.SaveAs(path)
        self.app.DisplayAlerts = True

    @activity
    def write_cell(self, column, row, value):
        """Write cell

        Write to a specific cell in the currently active workbook and active worksheet

        :parameter column: Column number (integer) to write
        :parameter row: Row number (integer) to write
        :parameter value: Value to write to specific cell

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Insert a text into the first cell
        >>> excel.write_cell(1,1, 'Hello World!')

        Keywords
            excel, cell, insert cell, insert data

        Icon
            las la-file-excel
        """
        self.workbook.ActiveSheet.Cells(row, column).Value = value

    @activity
    def read_cell(self, column, row):
        """Read cell

        Read a cell from the currently active workbook and active worksheet

        :parameter column: Column number (integer) to read
        :parameter row: Row number (integer) to read

        :return: Cell value

            :Example:
            
        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Insert a text into the first cell
        >>> excel.write_cell(1,1, 'Hello World!')
        >>> excel.read_cell(1,1)
        'Hello World!'

        Keywords
            excel, cell, read cell, read data

        Icon
            las la-file-excel
        """
        return self.workbook.ActiveSheet.Cells(row, column).Value

    @activity
    def write_range(self, range_, value):
        """Write range

        Write to a specific range in the currently active worksheet in the active workbook

        :parameter range_: Range to write to, e.g. "A1:D10"
        :parameter value: Value to write to range
        
            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Insert a text in every cell in this range
        >>> excel.write_range('A1:D5', 'Hello World!')

        Keywords
            excel, cell, write range, read data

        Icon
            las la-file-excel
        """
        self.workbook.ActiveSheet.Range(range_).Value = value

    @activity
    def read_range(self, range_):
        """Read range
        
        Read a range of cells from the currently active worksheet in the active workbook

        :parameter range_: Range to read from, e.g. "A1:D10"

        :return value: Values in param range
        
            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Insert a text in every cell in this range
        >>> excel.write_range('A1:D5', 'Hello World!')
        >>> # Read the same range
        >>> excel.read_range('A1:D5')
        [['Hello World', 'Hello World', 'Hello World', 'Hello World'], ...]

        Keywords
            excel, cell, read range, read data

        Icon
            las la-file-excel
        """
        return self.workbook.ActiveSheet.Range(range_).Value

    @activity
    def run_macro(self, name):
        """Run macro

        Run a macro by name from the currently active workbook

        :parameter name: Name of the macro to run. 

            :Example:

        >>> excel = Excel('excel_with_macro.xlsx')
        >>> # Run the macro
        >>> excel.run_macro('Macro1')

        Keywords
            excel, run macro, run vba

        Icon
            las la-file-excel
        """
        return self.app.Run(name)

    @activity
    def get_worksheet_names(self):
        """Get worksheet names

        Get names of all the worksheets in the currently active workbook

        :return: List is worksheet names

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add a worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Get all worksheet names
        >>> excel.get_worksheet_names()
        ['Sheet1', 'My Example Worksheet']

        Keywords
            excel, worksheet names, tab names

        Icon
            las la-file-excel
        """
        names = []

        for worksheet in self.workbook.Worksheets:
            names.append(worksheet.Name)

        return names

    @activity
    def get_table(self, name):
        """Get table

        Get table data from the currently active worksheet by name of the table

        :parameter name: List of table names

            :Example:
            
        >>> # Open Excel
        >>> excel = Excel()
        >>> # Create a table (Table1)
        >>> data = [
            {
                'Column A': 'Data Row 1 for A',
                'Column B': 'Data Row 1 for B',
                'Column C': 'Data Row 1 for C',
            },
            {
                'Column A': 'Data Row 2 for A',
                'Column B': 'Data Row 2 for B',
                'Column C': 'Data Row 2 for C',
            }]
        >>> excel.insert_data_as_table(data)
        >>> # Get the table
        >>> excel.get_table('Table1')
        [['Column A', 'Column B', 'Column C'], ['Row 1 A Data', 'Row 1 B Data', 'Row 1 C Data'], ...]

        Keywords
            excel, worksheet names, tab names

        Icon
            las la-file-excel
        """
        data = []

        for worksheet in self.workbook.Worksheets:
            for list_object in worksheet.ListObjects:
                if list_object.Name == name:
                    for row in list_object.DataBodyRange.Value:
                        data_row = {}
                        for i, column in enumerate(list_object.HeaderRowRange.Value[0]):
                            data_row[column] = row[i]
                        data.append(data_row)

        return data

    @activity
    def activate_range(self, range_):
        """Activate range

        Activate a particular range in the currently active workbook

        :parameter range_: Range to activate, e.g. "A1:D10"

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Activate a cell range
        >>> excel.activate_range('A1:D5')

        Keywords
            excel, activate range, make selection, select cells, select range

        Icon
            las la-file-excel
        """
        self.workbook.ActiveSheet.Range(range_).Select()

    @activity
    def activate_first_empty_cell_down(self):
        """Activate first empty cell down

        Activates the first empty cell going down

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Write some cells
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(1, 2, 'Filled')
        >>> excel.write_cell(1, 3, 'Filled')
        >>> # Activate the first empty cell going down, in this case cell A4 or (1,4)
        >>> excel.activate_first_empty_cell_down()

        Keywords
            excel, first empty cell, down

        Icon
            las la-file-excel
        """
        column = self.app.ActiveCell.Column
        row = self.app.ActiveCell.Row
        for cell in self.workbook.ActiveSheet.Columns(column).Cells:
            if not cell.Value and cell.Row > row:
                cell.Select()
                break

    @activity
    def activate_first_empty_cell_right(self):
        """Activate first empty cell right

        Activates the first empty cell going right

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Write some cells
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(1, 2, 'Filled')
        >>> excel.write_cell(1, 3, 'Filled')
        >>> # Activate the first empty cell going right, in this case cell B1 or (2,1)
        >>> excel.activate_first_empty_cell_right()

        Keywords
            excel, first empty cell, right

        Icon
            las la-file-excel
        """
        column = self.app.ActiveCell.Column
        row = self.app.ActiveCell.Row
        for cell in self.workbook.ActiveSheet.Rows(row).Cells:
            if not cell.Value and cell.Column > column:
                cell.Select()
                break

    @activity
    def activate_first_empty_cell_left(self):
        """Activate first empty cell left

        Activates the first empty cell going left

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(1, 2, 'Filled')
        >>> excel.write_cell(1, 3, 'Filled')
        >>> excel.activate_first_empty_cell_left()

        Keywords
            excel, first empty cell, left

        Icon
            las la-file-excel
        """
        column = self.app.ActiveCell.Column
        row = self.app.ActiveCell.Row

        for i in range(column):
            if column-i > 0:
                cell = self.workbook.ActiveSheet.Cells(row, column-i)
                if not cell.Value:
                    cell.Select()
                    break

    @activity
    def activate_first_empty_cell_up(self):
        """Activate first empty cell up

        Activates the first empty cell going up

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Write some cells
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(1, 2, 'Filled')
        >>> excel.write_cell(1, 3, 'Filled')
        >>> # Activate first empty cell
        >>> excel.activate_first_empty_cell_up()

        Keywords
            excel, first empty cell, up

        Icon
            las la-file-excel
        """
        column = self.app.ActiveCell.Column
        row = self.app.ActiveCell.Row

        for i in range(row):
            if row-i > 0:
                cell = self.workbook.ActiveSheet.Cells(row-i, column)
                if not cell.Value:
                    cell.Select()
                    break

    @activity
    def write_cell_formula(self, column, row, formula):
        """Write cell formula
        
        Write a formula to a particular cell

        :parameter column: Column number (integer) to write formula
        :parameter row: Row number (integer) to write formula
        :parameter value: Formula to write to specific cell e.g. "=10*RAND()"

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Write a formula to the first cell
        >>> excel.write_cell_formula(1, 1, '=1+1)

        Keywords
            excel, insert formula, insert calculation, insert calculated cell

        Icon
            las la-file-excel
        """
        self.workbook.ActiveSheet.Cells(row, column).Formula = formula

    @activity
    def read_cell_formula(self, column, row, formula):
        """Read cell formula

        Read the formula from a particular cell

        :parameter column: Column number (integer) to read formula
        :parameter row: Row number (integer) to read formula

        :return: Cell value

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Write a formula to the first cell
        >>> excel.write_cell_formula(1, 1, '=1+1)
        >>> # Read the cell
        >>> excel.read_cell_formula(1, 1)
        '=1+1'
        
        Keywords
            excel, read formula, read calculation

        Icon
            las la-file-excel
        """
        return self.workbook.ActiveSheet.Cells(row, column).Formula

    @activity
    def insert_empty_row(self, row):
        """Insert empty row

        Inserts an empty row to the currently active worksheet

        :parameter row: Row number (integer) where to insert empty row e.g 1
            
            :Example:

        >>> # Open Excel              
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(1, 2, 'Filled')
        >>> excel.write_cell(1, 3, 'Filled')
        >>> excel.insert_empty_row(2)
        
        Keywords
            excel, insert row, add row, empty row

        Icon
            las la-file-excel
        """
        row_range = 'A' + str(row)
        self.workbook.ActiveSheet.Range(row_range).EntireRow.Insert()

    @activity
    def insert_empty_column(self, column):
        """Insert empty column

        Inserts an empty column in the currently active worksheet. Existing columns will shift to the right.

        :parameter column: Column letter (string) where to insert empty column e.g. 'A'
            
            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(2, 2, 'Filled')
        >>> excel.write_cell(3, 3, 'Filled')
        >>> excel.insert_empty_column(2)
        
        Keywords
            excel, insert column, add column

        Icon
            las la-file-excel
        """
        column_range = str(column) + '1'
        self.workbook.ActiveSheet.Range(column_range).EntireColumn.Insert()

    @activity
    def delete_row(self, row):
        """Delete row in Excel

        Deletes a row from the currently active worksheet. Existing data will shift up.

        :parameter row: Row number (integer) where to delete row e.g 1

            :Example:

        >>> # Open Excel              
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(2, 2, 'Filled')
        >>> excel.write_cell(3, 3, 'Filled')
        >>> excel.delete_row(2)
        
        Keywords
            excel, delete row, remove row

        Icon
            las la-file-excel
        """
        row_range = 'A' + str(row)
        self.workbook.ActiveSheet.Range(row_range).EntireRow.Delete()

    @activity
    def delete_column(self, range_):
        """Delete column

        Delet a column from the currently active worksheet. Existing columns will shift to the left.

        :parameter column: Column letter (string) where to delete  column e.g. 'A'

            :Example:

        >>> # Open Excel              
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(2, 2, 'Filled')
        >>> excel.write_cell(3, 3, 'Filled')
        >>> excel.delete_column(2)
        
        Keywords
            excel, delete column, remove column

        Icon
            las la-file-excel
        """
        column_range = str(column) + '1'
        self.workbook.ActiveSheet.Range(column_range).EntireColumn.Delete()

    @activity
    def export_to_pdf(self, path=None):
        """Export to PDF

        :parameter path: Output path where PDF file will be exported to. Default path is home directory with filename 'pdf_export.pdf'.
        
            :Example:

        >>> # Open Excel              
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(2, 2, 'Filled')
        >>> excel.write_cell(3, 3, 'Filled')
        >>> excel.export_to_pdf('output.pdf')
        
        Keywords
            excel, save as pdf, export to pdf, export as pdf

        Icon
            las la-file-excel
        """
        if not path:
            import os
            path = os.path.join(os.path.expanduser("~"), 'pdf_export.pdf')

        self.workbook.ActiveSheet.ExportAsFixedFormat(0, path, 0, True, True)

    @activity
    def insert_data_as_table(self, data, range_='A1', table_style="TableStyleMedium2"):
        """Insert data as table
        
        Insert list of dictionaries as a table in Excel

        :parameter data: List of dictionaries to write as table
        :parameter range_: Range or startingpoint for table e.g. 'A1'
        
            :Example:
            
        >>> excel = Excel()
        >>> data = [
            {
                'Column A': 'Data Row 1 for A',
                'Column B': 'Data Row 1 for B',
                'Column C': 'Data Row 1 for C',
            },
            {
                'Column A': 'Data Row 2 for A',
                'Column B': 'Data Row 2 for B',
                'Column C': 'Data Row 2 for C',
            }
        >>> excel.insert_data_as_table(data)
        
        Keywords
            excel, insert data, insert table, create table

        Icon
            las la-file-excel
        """
        row = self.workbook.ActiveSheet.Range(range_).Row
        column = self.workbook.ActiveSheet.Range(range_).Column

        column_names = list(data[0].keys())
        data_values = [[d[key] for key in data[0].keys()] for d in data]

        values = [column_names] + data_values
        for i in range(len(values)):
            for j in range(len(values[0])):
                self.workbook.ActiveSheet.Cells(
                    row+i, column+j).Value = values[i][j]

        start_cell = self.workbook.ActiveSheet.Cells(row, column)
        end_cell = self.workbook.ActiveSheet.Cells(row+i, column+j)
        self.workbook.ActiveSheet.Range(start_cell, end_cell).Select()
        self.app.ActiveSheet.ListObjects.Add().TableStyle = table_style

    @activity
    def read_worksheet(self, name=None, headers=False):
        """Read worksheet

        Read data from a worksheet as a list of lists

        :parameter name: Optional name of worksheet to read. If no name is specified will take active sheet
        :parameter headers: Boolean to treat first row as headers. Default value is False

        :return: List of dictionaries with sheet data
        
            :Example:

        >>> # Open excel    
        >>> excel = Excel()
        >>> Write some cells
        >>> excel.write_cell(1, 1, 'A')
        >>> excel.write_cell(1, 2, 'B')
        >>> excel.write_cell(1, 3, 'C')
        >>> excel.read_worksheet()
        [['A'],['B'],['C']]
        
        Keywords
            excel, read worksheet, export data, read data

        Icon
            las la-file-excel
        """
        if name:
            self.activate_worksheet(name)

        data = self.workbook.ActiveSheet.UsedRange.Value

        if isinstance(data, str):
            return data

        # Remove empty columns and rows
        data = [list(x) for x in data if any(x)]
        transposed = list(map(list, zip(*data)))
        transposed = [row for row in transposed if any(row)]
        data = list(map(list, zip(*transposed)))

        if headers:
            header_row = data[0]
            data = data[1:]
            data = [{column: row[i]
                     for i, column in enumerate(header_row)} for row in data]

        return data

    @activity
    def quit(self):
        """Quit Excel

        This closes Excel, make sure to use :func: 'save' or 'save_as' if you would like to save before quitting.
        
            :Example:
            
        >>> # Open Excel  
        >>> excel = Excel()
        >>> # Quit Excel
        >>> excel.quit()
        
        Keywords
            excel, exit, quit, close

        Icon
            las la-file-excel
        """
        self.app.Application.Quit()

"""
Excel File
Icon: las la-file-excel
"""

class ExcelFile:
    @activity
    def __init__(self, file_path=None):
        """Read and Write xlsx files. 

        This activity can read, write and edit Excel (xlsx) files without the need of having Excel installed. 
        Note that, in contrary to working with the :func: 'Excel' activities, a file get saved directly after manipulation.

        :parameter file_path: Enter a path to open Excel with an existing Excel file. If no path is specified a 'workbook.xlsx' will be initialized in the home directory, this is the default value. If a workbook with the same name already exists the file will be overwritten.
        
            :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        
        Keywords
            excel, open, start, xlsx

        Icon
            las la-file-excel    

        """

        self.file_path = file_path
        self.sheet_name = None

        self.app = self._launch()

    def _launch(self):

        import openpyxl
        import os

        if self.file_path:
            if not os.path.exists(self.file_path):
                # self.book(self.file_path)
                self.book = openpyxl.load_workbook(self.file_path)
        else:
            path = os.path.join(os.path.expanduser("~"), 'workbook.xlsx')
            self.book = openpyxl.load_workbook(path)
            self.file_path = path

    @activity
    def activate_worksheet(self, name):
        """Activate worksheet

        Activate a worksheet. By default the first worksheet is activated.

        :parameter name: Name of the worksheet to activate.        
        
            :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Add some worksheets
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> excel_file.add_worksheet('Another Worksheet')
        >>> # Activate a worksheet
        >>> excel_file.active_worksheet('My Example Worksheet')
        
        Keywords
            excel, activate tab, activate worksheet

        Icon
            las la-file-excel
        """

        self.sheet_name = name

    @activity
    def save_as(self, path):
        """Save as

        :parameter path: Path where workbook will be saved
        
            :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Ad a worksheet
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> # Save the Excel file
        >>> excel_file.save_as('output.xlsx')
        >>> # We can also save it in the home directory by using
        >>> excel_file.save_as( home_path('output.xlsx') )
        
        Keywords
            excel, save as, export, save

        Icon
            las la-file-excel
        """
        self.book.save(path)

    @activity
    def write_cell(self, column, row, value, auto_save=True):
        """Write cell

        :parameter column: Column number (integer) to write
        :parameter row: Row number (integer) to write
        :parameter value: Value to write to specific cell
        :parameter auto_save: Save document after performing activity. Default value is True
        
            :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Add a worksheet
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> excel_file.write_cell(1, 1, 'Filled!')
        
        Keywords
            excel, write cell, insert data

        Icon
            las la-file-excel
        """
        if self.sheet_name:
            sheet = book[self.sheet_name]
        else:
            sheet = book.active

        sheet.cell(row=row, column=column).value = value

        if auto_save:
            self.book.save(self.file_path)

    @activity
    def read_cell(self, column, row):
        """ Read cell

        :parameter column: Column number (integer) to read
        :parameter row: Row number (integer) to read

        :return: Cell value

           :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Add a worksheet
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> # Write the first cell
        >>> excel_file.write_cell(1, 1, 'Filled!')
        >>> # Read the first cell
        >>> excel_file.read_cell(1, 1)
        'Filled!'
        
        Keywords
            excel, read cell, read

        Icon
            las la-file-excel

        """
        if self.sheet_name:
            sheet = self.book[self.sheet_name]
        else:
            sheet = self.book.active

        return sheet.cell(row=row, column=column).value

    @activity
    def add_worksheet(self, name, auto_save=True):
        """Add worksheet

        :parameter name: Name of the worksheet to add
        :parameter auto_save: Save document after performing activity. Default value is True

           :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Add a worksheet
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> # List all the worksheets
        >>> excel.get_worksheet_names()

        
        Keywords
            excel, add worksheet, worksheet

        Icon
            las la-file-excel
        """

        self.book.create_sheet(name)
        if auto_save:
            self.book.save(self.file_path)

    @activity
    def get_worksheet_names(self):
        """Get worksheet names

        :return: List of worksheet names

           :Example:
            
        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Add some worksheets
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> excel_file.add_worksheet('Another Worksheet')
        >>> # Get the worksheet names
        >>> excel_file.get_worksheet_names()
        ['My Example Worksheet', 'Another Worksheet']
        
        Keywords
            excel, worksheet names, worksheet,

        Icon
            las la-file-excel

        """

        return self.book.sheetnames


"""
PowerPoint Application
icon: las la-file-powerpoint
"""


class PowerPoint:
    @activity
    def __init__(self, visible=True, path=None, add_slide=True):
        """Start PowerPoint Application

        For this activity to work, PowerPoint needs to be installed on the system.

        :parameter visible: Show PowerPoint in the foreground if True or hide if False, defaults to True.
        :parameter path: Enter a path to open an existing PowerPoint presentation. If no path is specified a new presentation will be initialized, this is the default value.
        :parameter add_slide: Add an initial empty slide when creating new PowerPointfile, this prevents errors since most manipulations require a non-empty presentation. Default value is True
        
            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()

        Keywords
            powerpoint, ppt

        Icon
            las la-file-powerpoint
        
        """
        self.app = self._launch(path)
        self.app.Visible = visible

    def _launch(self, path):
        """Utility function to create the Excel application scope object

        :return: Application object (win32com.client)
        """
        try:
            import win32com.client

            app = win32com.client.gencache.EnsureDispatch(
                "PowerPoint.Application")

        except:
            raise Exception(
                "Could not launch PowerPoint, do you have Microsoft Office installed on Windows?")

        if path:
            return app.Presentations.Open(file_path)
        else:
            return app.Presentations.Add()

    @activity
    def save_as(self, path=None):
        """Save PowerPoint
        
        Save PowerPoint Slidedeck

        :parameter path: Save the PowerPoint presentation. Default value is the home directory and filename 'presentation.pptx'

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add a first slide
        >>> powerpoint.add_slide()
        >>> # Save the PowerPoint presentation
        >>> powerpoint.save()

        Keywords
            powerpoint, ppt, save, save as, save powerpoint

        Icon
            las la-file-powerpoint

        """
        if not path:
            import os
            os.path.join(os.path.expanduser("~"), 'presentation.pptx')

        return self.app.SaveAs(path)

    @activity
    def quit(self):
        """Close PowerPoint Application

        Close PowerPoint

        :parameter index: Index where the slide should be inserted. Default value is as final slide.
        :parmeter type: Type of the slide to be added. Supports following types: blank, chart, text, title and picture.

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Close PowerPoint
        >>> powerpoint.quit()


        Keywords
            powerpoint, ppt, quit, exit

        Icon
            las la-file-powerpoint

        """
        return self.app.Application.Quit()

    @activity
    def add_slide(self, index=None, type='blank'):
        """Add PowerPoint Slides

        Adds slides to a presentation

        :parameter index: Index where the slide should be inserted. Default value is as final slide.
        :parmeter type: Type of the slide to be added. Supports following types: blank, chart, text, title and picture.

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add a first slide
        >>> powerpoint.add_slide()


        Keywords
            powerpoint, ppt, add, add slide  powerpoint, slides

        Icon
            las la-file-powerpoint

        """
        if type == 'blank':
            type_id = 12
        if type == 'chart':
            type_id = 8
        if type == 'text':
            type_id = 2
        if type == 'title':
            type_id = 1
        if type == 'picture':
            type_id = 36

        if not index:
            index = self.app.Slides.Count + 1

        return self.app.Slides.Add(index, type_id)

    @activity
    def number_of_slides(self):
        """Slide count
        
        :return: The number of slides

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add some slides
        >>> powerpoint.add_slide()
        >>> powerpoint.add_slide()
        >>> # Show number of slides
        >>> powerpoint.number_of_slides()

        Keywords
            powerpoint, ppt, slide count, number of slides

        Icon
            las la-file-powerpoint
        """
        return self.app.Slides.Count

    @activity
    def add_text(self, text, index=None, font_size=48, font_name=None, bold=False, margin_bottom=100, margin_left=100, margin_right=100, margin_top=100):
        """Text to slide
        
        Add text to a slide

        :parameter index: Slide index to add text. If none is specified, a new slide will be added as final slide
        :parmeter text: Text to be added
        :parameter font_size: Fontsize, default value is 48
        :parameter font_name: Fontname, if not specified will take default PowerPoint font
        :parameter bold: Toggle bold with True or False, default value is False
        :parameter margin_bottom: Margin from the bottom in pixels, default value is 100 pixels
        :parameter margin_left: Margin from the left in pixels, default value is 100 pixels
        :parameter margin_right: Margin from the right in pixels, default value is 100 pixels
        :parameter margin_top: Margin from the top in pixels, default value is 100 pixels

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add slide with text
        >>> powerpoint.add_text(text='Sample Text')


        Keywords
            powerpoint, ppt, text, add text, slides
        Icon
            las la-file-powerpoint

        """

        if not index:
            index = self.app.Slides.Count + 1
            self.app.Slides.Add(index, 12)
        text_box = self.app.Slides(index).Shapes.AddTextbox(
            1, 100, 100, 200, 50).TextFrame.TextRange
        text_box.Text = text
        text_box.Font.Size = font_size
        if font_name:
            text_box.Font.Name = font_name
        text_box.Font.Bold = bold

    @activity
    def delete_slide(self, index=None):
        """Delete slide

        :parameter index: Slide index to be deleted. If none is specified, last slide will be deleted

            :Example:

        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add some slides
        >>> powerpoint.add_slide()
        >>> powerpoint.add_slide()
        >>> # Delete last slide
        >>> powerpoint.delete_slide()

        Keywords
            powerpoint, ppt, delete, delete slide

        Icon
            las la-file-powerpoint

        """
        if not index:
            index = self.app.Slides.Count

        return self.app.Slides(index).Delete()

    @activity
    def replace_text(self, placeholder_text, replacement_text):
        """Replace all occurences of text in PowerPoint slides

        Can be used for example to replace arbitrary placeholder value in a PowerPoint. 
        For example when using a template slidedeck, using 'XXXX' as a placeholder.
        Take note that all strings are case sensitive.

        :parameter placeholder_text: Placeholder value (string) in the PowerPoint, this will be replaced, e.g. 'Company Name'
        :parameter replacement_text: Text (string) to replace the placeholder values with. It is recommended to make this unique in your PowerPoint to avoid wrongful replacement, e.g. 'XXXX_placeholder_XXX'

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add some slides with text
        >>> powerpoint.add_text(text='Hello, my name is placeholder')
        >>> # Change 'placeholder' to the word 'robot
        >>> powerpoint.replace_text(placeholder_text = 'placeholder', replacement_text ='robot')

        Keywords
            powerpoint, ppt, replace, placeholder

        Icon
            las la-file-powerpoint

        """
        for slide in self.app.Slides:
            for shape in slide.Shapes:
                shape.TextFrame.TextRange.Text = shape.TextFrame.TextRange.Text.replace(placeholder_text, replacement_text)

    @activity
    def export_to_pdf(self, path=None):
        """PowerPoint to PDF
        
        Export PowerPoint presentation to PDF file

        :parameter path: Output path where PDF file will be exported to. Default path is home directory with filename 'pdf_export.pdf'.

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add some slides with text
        >>> powerpoint.add_text(text='Robots are cool')
        >>> # Export to pdf
        >>> powerpoint.export_to_pdf()

        Keywords
            powerpoint, ppt, export, pdf

        Icon
            las la-file-powerpoint

        """

        if self.app.Slides.Count == 0:
            raise Exception(
                'Please add a slide first bedore exporting the presentation.')

        if not path:
            import os
            path = os.path.join(os.path.expanduser("~"), 'pdf_export.pdf')

        return self.app.ExportAsFixedFormat2(path, 2, PrintRange=None)

    @activity
    def export_slides_to_images(self, path=None, type='png'):
        """Slides to images
        
        Export PowerPoint slides to seperate image files

        :parameter path: Output path where image files will be exported to. Default path is home directory.
        :parameter type: Output type of the images, supports 'png' and 'jpg' with 'png' as default value

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add some slides with text
        >>> powerpoint.add_text(text='Robots are cool')
        >>> powerpoint.add_text(text='Humans are cooler')
        >>> # Export slides to images
        >>> powerpoint.export_slides_to_images()

        Keywords
            powerpoint, ppt, export, png, image, slides to image

        Icon
            las la-file-powerpoint

        """

        if self.app.Slides.Count == 0:
            raise Exception(
                'Please add a slide first bedore exporting the presentation.')

        if not path:
            import os
            path = os.path.expanduser("~")

        return self.app.Export(path, 'png')


"""
Office 365
Icon: las la-cloud
"""


@activity
def send_email_with_outlook365(client_id, client_secret, to_email, subject='', body=''):
    """Send email Office Outlook 365

    :parameter client_id: Client id for office 365 account
    :parameter client_secret: Client secret for office 365 account
    :parameter to_email: E-mail to send to
    :parameter subject: Optional subject
    :parameter body: Optional body of the email

        :Example:

    >>> # Send email to 'robot@automagica.com'
    >>> send_email_with_outlook365('SampleClientID', 'SampleClientSecret', 'robot@automagica.com')

    Keywords
        mail, office 365, outlook, email, e-mail

    Icon
        las la-envelope
    """
    from O365 import Account

    credentials = (client_id, client_secret)

    account = Account(credentials)
    m = account.new_message()
    m.to.add(to_email)
    m.subject = subject
    m.body = body
    m.send()


"""
Salesforce
Icon: lab la-salesforce
"""

@activity
def salesforce_api_call(action, key, parameters={}, method='get', data={}):
    """Salesforce API

    Activity to make calls to Salesforce REST API.

    :parameter action: Action (the URL)
    :parameter key: Authorisation key 
    :parameter parameters: URL params
    :parameter method: Method (get, post or patch)
    :parameter data: Data for POST/PATCH.

    :return: API data

        :Example:

    >>> spf_api_call('action', 'key', 'parameters')
    Response

    Keywords
        salesforce

    Icon
        lab la-salesforce

    """
    headers = {
        'Content-type': 'application/json',
        'Accept-Encoding': 'gzip',
        'Authorization': 'Bearer ' + key
    }
    if method == 'get':
        r = requests.request(method, instance_url+action,
                             headers=headers, params=parameters, timeout=30)
    elif method in ['post', 'patch']:
        r = requests.request(method, instance_url+action,
                             headers=headers, json=data, params=parameters, timeout=10)
    else:
        raise ValueError('Method should be get or post or patch.')
    print('Debug: API %s call: %s' % (method, r.url))
    if r.status_code < 300:
        if method == 'patch':
            return None
        else:
            return r.json()
    else:
        raise Exception('API error when calling %s : %s' % (r.url, r.content))


"""
E-mail (SMTP)
Icon: las la-at
"""


@activity
def send_mail_smtp(smtp_host, smtp_user, smtp_password, to_address, subject="", message="", port=587):
    """Mail with SMTP

    This function lets you send emails with an e-mail address. 

    :parameter smpt_host: The host of your e-mail account. 
    :parameter smpt_user: The password of your e-mail account
    :parameter smpt_password: The password of your e-mail account
    :parameter to_address: The destination is the receiving mail address. 
    :parameter subject: The subject 
    :parameter message: The body of the mail
    :parameter port: The port variable is standard 587. In most cases this argument can be ignored, but in some cases it needs to be changed to 465.

        :Example:

    >>> send_mail_smpt('robot@automagica.com', 'SampleUser', 'SamplePassword', 'robotfriend@automagica.com')

    Keywords
        mail, e-mail, email smpt

    Icon
        las la-mail-bulk

    """
    BODY = "\r\n".join(
        [
            "To: %s" % destination,
            "From: %s" % user,
            "Subject: %s" % subject,
            "",
            message,
        ]
    )
    smtpObj = smtplib.SMTP(host, port)
    smtpObj.ehlo()
    smtpObj.starttls()
    smtpObj.login(user, password)
    smtpObj.sendmail(user, destination, BODY)
    smtpObj.quit()


"""
Windows OS
icon: lab la-windows
"""


@activity
def set_user_password(username, password):
    """Set Windows password

    Sets the password for a Windows user.

    :parameter username: Username
    :parameter password: New password

        :Example:

    >>> set_user_password('SampleUsername', 'SamplePassword')

    Keywords
        windows, user, password, account

    Icon
        las la-passport

    """
    from win32com import adsi

    user = adsi.ADsGetObject("WinNT://localhost/%s,user" % username)
    user.SetPassword(password)


@activity
def validate_user_password(username, password):
    """Check Windows password

    Validates a Windows user password if it is correct

    :parameter username: Username
    :parameter password: New password

    :return: True if the password is correct

        :Example:

    >>> validate_user_password('SampleUsername', 'SamplePassword')
    False

    Keywords
        windows, user, password, account

    Icon
        las la-passport

    """
    from win32security import LogonUser
    from win32con import LOGON32_LOGON_INTERACTIVE, LOGON32_PROVIDER_DEFAULT

    try:
        LogonUser(
            username,
            None,
            password,
            LOGON32_LOGON_INTERACTIVE,
            LOGON32_PROVIDER_DEFAULT,
        )
    except:
        return False
    return True


@activity
def lock_windows():
    """Lock Windows

    Locks Windows requiring login to continue.

        :Example:

    >>> lock_windows()

    Keywords
        windows, user, password, account, lock, freeze, hibernate, sleep, lockescreen

    Icon
        las la-user-lock

    """
    import ctypes
    ctypes.windll.user32.LockWorkStation()


@activity
def is_logged_in():
    """Check if Windows logged in

    Checks if the current user is logged in and not on the lockscreen. Most automations do not work properly when the desktop is locked.

    :return: True if the user is logged in, False if not

        :Example:

    >>> is_logged_in()
    True

    Keywords
        windows, login, logged in, lockscreen, user, password, account, lock, freeze, hibernate, sleep

    Icon
        lar la-user
    """
    import subprocess

    output = subprocess.check_output("TASKLIST")

    if "LogonUI.exe" in str(output):
        return False
    else:
        return True


@activity
def is_desktop_locked():
    """Check if Windows is locked

    Checks if the current user is locked out and on the lockscreen. Most automations do not work properly when the desktop is locked.

    :return: True when the lockscreen is active, False if not.

        :Example:

    >>> desktop_locked()
    True

    Keywords
        windows, login, logged in, lockscreen, user, password, account, lock, locked, freeze, hibernate, sleep

    Icon
        las la-user

    """
    return not is_logged_in()


@activity
def get_username():
    """Get Windows username

    Get current logged in user's username

        :Example:

    >>> get_username()
    'Automagica'

    Keywords
        windows, login, logged in, lockscreen, user, password, account, lock, locked, freeze, hibernate, sleep

    Icon
        las la-user

    """
    import getpass
    return getpass.getuser()


@activity
def set_to_clipboard(text):
    """Set clipboard

    Set any text to the Windows clipboard. 

    :parameter text: Text to put in the clipboard

        :Example:

    >>> # Create some sample text
    >>> sample_text = 'A robots favourite food must be computer chips'
    >>> # Set to clipboard
    >>> set_to_clipboard(sample_text)
    >>> # Print the clipboard to verify
    >>> print( get_from_clipboard() )

    Keywords
        copy, clipboard, clip board, ctrl c, ctrl v, paste

    Icon
        las la-clipboard-check
    """
    import win32clipboard

    win32clipboard.OpenClipboard()
    win32clipboard.EmptyClipboard()
    win32clipboard.SetClipboardText(text, win32clipboard.CF_UNICODETEXT)
    win32clipboard.CloseClipboard()


@activity
def get_from_clipboard():
    """Get clipboard

    Get the text currently in the Windows clipboard

    :return: Text currently in the clipboard

        :Example:

    >>> # Create some sample text
    >>> sample_text = 'A robots favourite food must be computer chips'
    >>> # Set to clipboard
    >>> set_to_clipboard(sample_text)
    >>> # Get the clipboard to verify
    >>> get_from_clipboard()
    'A robots favourite food must be computer chips'

    Keywords
        copy, clipboard, clip board, ctrl c, ctrl v, paste

    Icon
        las la-clipboard-list

    """
    import win32clipboard

    win32clipboard.OpenClipboard()
    try:
        data = str(win32clipboard.GetClipboardData(
            win32clipboard.CF_UNICODETEXT))
        return data

    except:
        return None

    finally:
        win32clipboard.CloseClipboard()


@activity
def clear_clipboard():
    """Empty clipboard

    Empty text from clipboard. Getting clipboard data after this should return in None

        :Example:

    >>> # Create some sample text
    >>> sample_text = 'A robots favourite food must be computer chips'
    >>> # Set to clipboard
    >>> set_to_clipboard(sample_text)
    >>> # Clear the clipboard
    >>> clear_clipboard()
    >>> # Get clipboard contents to verify
    >>> print( get_clipboard() )
    None

    Keywords
        copy, clipboard, clip board, ctrl c, ctrl v, paste

    Icon
        las la-clipboard
    """
    from ctypes import windll

    if windll.user32.OpenClipboard(None):
        windll.user32.EmptyClipboard()
        windll.user32.CloseClipboard()
    return


@activity
def run_vbs_script(script_path, parameters=[]):
    """Run VBSscript

    Run a VBScript file

    :parameter script_path: Path to the .vbs-file
    :parameter parameters: Additional arguments to pass to the VBScript

        :Example:

    >>> # Run a VBS script
    >>> run_vbs_script('Samplescript.vbs')

    Keywords
        vbs, VBScript

    Icon
        las la-cogs
    """
    import subprocess

    subprocess.call(["cscript.exe", script_path] + parameters)


@activity
def beep(frequency=1000, duration=500):
    """Beep

    Make a beeping sound. Make sure your volume is up and you have hardware connected.

    :parameter frequency: Integer to specify frequency (Hz), default value is 1000 Hz
    :parameter duration: Integer to specify duration of beep in miliseconds (ms), default value is 500 ms.

    :return: Sound

        :Example:

    >>> beep()

    Keywords
        beep, sound, noise, speaker, alert

    Icon
        las la-volume-up

    """
    import winsound

    winsound.Beep(frequency, duration)


"""
Text-to-Speech
Icon: las la-volume-up
"""


@activity
def speak(text, speed=None):
    """Speak

    Use the Text-To-Speech engine available on your system to read text

    :parameter text: The text which should be said
    :parameter speed: Multiplication factor for the speed at which the text should be pronounced. 

    :return: Spoken text

        :Example:

    >>> # Read the following text out loud
    >>> speak('How do robots eat guacamole?')
    >>> speak('With microchips!')

    Keywords
        sound, speech, text, speech to text, speech-to-text, translate, read, read out loud

    Icon
        las la-microphone-alt


    """
    import pyttsx3

    engine = pyttsx3.init()

    if speed:
        default_rate = engine.getProperty("rate")
        engine.setProperty("rate", speed * default_rate)

    engine.say(text)
    engine.runAndWait()


"""
Active Directory
Icon: las la-user
"""


class ActiveDirectory():

    @activity
    def __init__(self, ldap_server=None, username=None, password=None):
        """AD interface
        
        Interface to Windows Active Directory through ADSI

        Activity to connect the ADSI interface to Microsoft Active Directory.
        Connects to the AD domain to which the machine is joined by default.

            :Example:

        >>> ad = ActiveDirectory()

        Keywords
            AD, active directory, activedirectory

        Icon
            las la-audio-description

        """
        import pyad

        self.pyad = pyad

        if ldap_server:
            self.pyad.set_defaults(ldap_server=ldap_server)

        if username:
            self.pyad.set_defaults(username=username)

        if password:
            self.pyad.set_defaults(password=password)

    @activity
    def get_object_by_distinguished_name(self, distinguished_name):
        """Get AD object by name
        
        Interface to Windows Active Directory through ADSI

        Activity to connect the ADSI interface to Microsoft Active Directory.
        Connects to the AD domain to which the machine is joined by default.

            :Example:

        >>> ad = ActiveDirectory()
        >>> ad.get_object_by_distinguished_name('SampleDN')

        Keywords
            AD, active directory, activedirectory

        Icon
            las la-audio-description

        """
        return self.pyad.from_dn(distinguished_name)


"""
Utilities
icon: las la-toolbox
"""


@activity
def home_path(subdir=None):
    """Get user home path

    Returns the current user's home path

    :parameter filename: Optional filename to add to the path. Can also be a subdirectory

    :return: Path to the current user's home folder

        :Example:

    >>> # Home_path without arguments will return the home path
    >>> print( home_path() )
    >>> # When looking for a file in the home path, we can specify it
    >>> # First make a sample textfile
    >>> make_textfile()
    >>> # Refer to it
    >>> home_path('generated_textfile.txt')
    'C:\\Users\\<username>\\generated_textfile.txt'

    Keywords
        home, home path, homepath, home directory, homedir

    Icon
        las la-home

    """
    import os

    if subdir:
        return os.path.join(os.path.expanduser("~"), subdir)
    return os.path.expanduser("~")


@activity
def desktop_path(subdir=None):
    """Get desktop path

    Returns the current user's desktop path

    :parameter filename: Optional filename to add to the path. Can also be a subdirectory

    :return: Path to the current user's desktop folder

        :Example:

    >>> # Desktop_path without arguments will return the home path
    >>> print( desktop_path() )
    >>> # When looking for a file on the desktop, we can specify it
    >>> # First make a sample textfile
    >>> make_textfile()
    >>> # Refer to it
    >>> desktop_path('generated_textfile.txt')
    'C:\\Users\\<username>\\Desktop\\generated_textfile.txt'

    Keywords
        desktop, desktop path, desktoppath, desktop directory, desktopdir

    Icon
        lar la-desktop
    """
    import os

    if subdir:
        return os.path.join(os.path.join(os.path.expanduser("~"), "Desktop"), subdir)
    return os.path.join(os.path.expanduser("~"), "Desktop")


@activity
def open_file(path):
    """Open file

    Opens file with default programs

    :parameter path: Path to file. 

    :return: Path to file

        :Example:

    >>> # Make textfile
    >>> testfile = make_textfile()
    >>> # Open the file
    >>> open_file(testfile)

    Keywords
        file, open, open file, show, reveal, explorer, run, start

    Icon
        lar la-file

    """

    import os
    os.startfile(path)

    return path


@activity
def set_wallpaper(image_path):
    """Set wallpaper

    Set Windows desktop wallpaper with the the specified image

    :parameter image_path: Path to the image. This image will be set as desktop wallpaper

        :Example:

    >>> # Caution: this example will change your wallpaper
    >>> # Take a screenshot of current screen
    >>> screenshot = take_screenshot()
    >>> # Flip it hozirontally for fun
    >>> mirror_image_horizontally(screenshot)
    >>> # Set flipped image as wallpaper
    >>> set_wallpaper(screenshot)

    Keywords
        desktop, desktop path, desktoppath, desktop directory, desktopdir

    Icon
        las la-desktop

    """
    import ctypes
    ctypes.windll.user32.SystemParametersInfoW(20, 0, image_path, 0)


@activity
def download_file_from_url(url, filename=None, path=None):
    """Download file from a URL

    :parameter url: Source URL to download file from
    :parameter filename: 
    :parameter path: Target path. If no path is given will download to the home directory

    :return: Target path as string

        :Example:

    >>> # Download robot picture from the wikipedia robot page
    >>> picture_url = 'https://upload.wikimedia.org/wikipedia/commons/thumb/6/6c/Atlas_from_boston_dynamics.jpg/220px-Atlas_from_boston_dynamics.jpg'
    >>> download_file_from_url(url = picture_url, filename = 'robot.jpg')
    'C:\\Users\\<username>\\robot.jpg'

    Keywords
        download, download url, save, request

    Icon
        las la-cloud-download-alt
    """
    import requests
    import re
    import os

    if not filename:
        base_path, filename = os.path.split(url)
    if not path:
        path = os.path.join(os.path.expanduser("~"), filename)

    r = requests.get(url, stream=True)

    if r.status_code == 200:
        with open(path, 'wb') as f:
            f.write(r.content)

        return path

    else:
        raise Exception('Could not download file from {}'.format(url))


"""
Trello
Icon: lab la-trello
"""


@activity
def add_trello_card(title="My card", description="My description", board_name="My board", list_name="My list", api_key="", api_secret="", token="", token_secret="any"):
    """Add Trello Card

    Add a card to the Trello board. For this you need a Trello API key, secret and token. 


    :parameter title: Title of Trello card
    :parameter description: Description of Trello card
    :parameter board_name: Name of the Trello board
    :parameter api_key: Trello API key
    :parameter api_secret: Trello API secret
    :parameter token: Trello token
    :parameter token_secret: Token secret can be any string, but should be altered for security purposes.

        :Example:

    >>> add_trello_card(title='ExampleTitle', description='ExampleDescription', api_key='SampleKey', api_secret='ApiSecret', token='SampleToken')

    Keywords
        trello

    Icon
        lab la-trello

    """
    from trello import TrelloClient

    client = TrelloClient(
        api_key=api_key,
        api_secret=api_secret,
        token=token,
        token_secret=token_secret,
    )

    trello_boards = client.list_boards()
    for trello_board in trello_boards:
        if trello_board.name == board_name:
            target_board = trello_board
            break

    trello_lists = target_board.all_lists()
    for trello_list in trello_lists:
        if trello_list.name == list_name:
            target_list = trello_list
            break

    target_list.add_card(title, desc=description)


"""
System
Icon: las la-laptop
"""


@activity
def rename_file(input_path, new_name=None):
    """Rename a file

    This activity will rename a file. If the the desired name already exists in the folder file will not be renamed.

    :parameter path: Full path to file that will be renamed
    :parameter new_name: New name of the file e.g. 'newfile.txt'. By default file will be renamed to original folder name with '_renamed' added to the folder name.

    :return: Path to renamed file as a string. None if folder could not be renamed.

        :Example:

    >>> # Make new textfile in home directory
    >>> textfile = make_textfile()
    >>> # Rename the file
    >>> rename_file(textfile)
    C:\\Users\\<username>\\generated_textfile_renamed.txt'

    Keywords
        file, rename, rename file, organise file, files, file manipulation, explorer, nautilus

    Icon
        las la-file-contract
    """
    import os

    if not os.path.isfile(input_path):
        return None

    if not new_name:
        base, file_extension = os.path.splitext(input_path)
        new_path = base + '_renamed' + file_extension
    else:
        base_path, filename = os.path.split(input_path)
        new_path = os.path.join(base_path, new_name)

    if os.path.isfile(new_path):
        return None

    os.rename(input_path, new_path)
    return new_path


@activity
def move_file(from_path, to_path):
    """Move a file

    If the new location already contains a file with the same name, a random 4 character uid will be added in front of the name before the file is moved.

    :parameter old_path: Full path to the file that will be moved
    :parameter new_location: Path to the folder where file will be moved to

    :return: Path to renamed file as a string. None if folder could not be moved.

        :Example:

    >>> # Make new textfile in home directory
    >>> textfile = make_textfile()
    >>> # Make a folder to move the file to
    >>> new_folder = create_folder()
    >>> # Move textfile to the folder
    >>> move_file(textfile, new_folder)

    Keywords
        file, move, move file, organise file, files, file manipulation, explorer, nautilus

    Icon
        las la-file-export

    """
    import uuid
    import os
    import shutil

    if not os.path.isfile(from_path):
        return None

    if os.path.isfile(to_path):
        base, file_extension = os.path.splitext(to_path)
        to_path = base + str(uuid4())[:4] + file_extension

    shutil.move(from_path, to_path)
    return to_path


@activity
def remove_file(path):
    """Remove a file

    :parameter path: Full path to the file that will be deleted.

    :return: Path to removed file as a string.

        :Example:

    >>> # Make new textfile in home directory
    >>> textfile = make_textfile()
    >>> # Remove the file
    >>> remove_file(textfile)

    Keywords
        file, delete, erase, delete file, organise file, files, file manipulation, explorer, nautilus

    Icon
        las la-trash
    """

    import os
    if os.path.isfile(path):
        os.remove(path)
    return path


@activity
def file_exists(path):
    """Check if file exists

    This function checks whether the file with the given path exists.

    :parameter path: Full path to the file to check.

    return: True or False (boolean)

        :Example:

    >>> # Make new textfile in home directory
    >>> textfile = make_textfile()
    >>> # Check if file exists
    >>> file_exists(textfile)
    True

    Keywords
        file, exists, files, file manipulation, explorer, nautilus

    Icon
        las la-tasks
    """
    import os
    return os.path.isfile(path)


@activity
def wait_file_exists(path, timeout=60):
    """Wait until a file exists.

    Not that this activity is blocking and will keep the system waiting.

    :parameter path: Full path to file.
    :parameter timeout: Maximum time in seconds to wait before continuing. Default value is 60 seconds.

        :Example:

    >>> # Make new textfile in home directory
    >>> textfile = make_textfile()
    >>> # Wait untile file exists # Should pass immediatly
    >>> wait_file_exists(textfile)

    Keywords
        file, wait, wait till exists, files, file manipulation, explorer, nautilus

    Icon
        las la-list-alt

    """
    from time import sleep

    while not os.path.exists(path):
        sleep(1)
    return

    for _ in range(timeout):
        if os.path.exists(path):
            break
            sleep(1)


@activity
def write_list_to_file(list_to_write, file_path):
    """List to .txt
    
    Writes a list to a  text (.txt) file. 
    Every element of the entered list is written on a new line of the text file.

    :parameter list_to_write: List to write to .txt file
    :parameter path: Path to the text-file. 

        :Example:
    
    >>> # Make a list to write
    >>> robot_names = ['WALL-E', 'Terminator', 'R2D2']
    >>> # Create a new text file
    >>> textfile = make_textfile()
    >>> write_list_to_file(robot_names, textfile)
    >>> # Open the file for illustration
    >>> open_file(textfile)

    Keywords
        list, text, txt, list to file, write list, write

    Icon
        las la-list

    """
    with open(file_path, "w") as filehandle:
        filehandle.writelines("%s\n" % place for place in list_to_write)
    return


@activity
def read_list_from_txt(file_path):
    """Read .txt file

    This activity writes the content of a .txt file to a list and returns that list. 
    Every new line from the .txt file becomes a new element of the list. The activity will 
    not work if the entered path is not attached to a .txt file.

    :parameter path: Path to the .txt file

    :return: List with contents of specified .txt file

        :Example:
    
    >>> # Make a list to write
    >>> robot_names = ['WALL-E', 'Terminator', 'R2D2']
    >>> # Create a new text file
    >>> textfile = make_textfile()
    >>> write_list_to_file(robot_names, textfile)
    >>> # Read list from file
    >>> read_list_from_txt(textfile)
    ['WALL-E', 'Terminator', 'R2D2']

    Keywords
        list, text, txt, list to file, write list, read, read txt, read text

    Icon
        las la-th-list

    """
    written_list = []
    with open(file_path, "r") as filehandle:
        filecontents = filehandle.readlines()
        for line in filecontents:
            current_place = line[:-1]
            written_list.append(current_place)
    return written_list


@activity
def append_line(text, file_path):
    """Append to .txt
    
    Append a text line to a file and creates the file if it does not exist yet.

    :parameter text: The text line to write to the end of the file
    :parameter file_path: Path to the file to write to

        :Example:
    
    >>> # Create a new text file
    >>> textfile = make_textfile()
    >>> # Append a few lines to the file
    >>> append_line('Line 1', textfile)
    >>> append_line('Line 2', textfile)
    >>> append_line('Line 3', textfile)
    >>> # Open the file for illustration
    >>> open_file(textfile)

    Keywords
        list, text, txt, list to file, write list, read, write txt, append text, append line, append, add to file, add

    Icon
        las la-tasks
    """
    
    import os

    if not os.path.isfile(file_path):
        with open(file_path, "a"):
            os.utime(file_path, None)

    with open(file_path, "a") as f:
        f.write("\n" + text)


@activity
def make_textfile(text='Sample text', output_path=None):
    """Make text file

    Initialize text file

    :parameter text: The text line to write to the end of the file. Default text is 'Sample text'
    :parameter output_path: Ouput path. Will write to home directory

    :return: Path as string

        :Example:

    >>> # Create a new text file
    >>> textfile = make_textfile()
    C:\\Users\\<username>\\generated_textfile.txt'

    Keywords
        make textfile, textfile, testfile, exampel file, make file, make, new file, new textfile, txt, new txt

    Icon
        las la-file-alt

    """

    # Set to user home if no path specified
    import os
    if not output_path:
        output_path = os.path.join(
            os.path.expanduser("~"), "generated_textfile.txt")
    with open(output_path, "w") as file:
        file.write(text)

    return output_path


@activity
def copy_file(old_path, new_path=None):
    """Copy a file

    Copies a file from one place to another.
    If the new location already contains a file with the same name, a random 4 character uid is added to the name.

    :parameter old_path: Full path to the source location of the folder
    :parameter new_path: Optional full path to the destination location of the folder. If not specified file will be copied to the same location with a random 8 character uid is added to the name.
    
    :return: New path as string

        :Example:

    >>> # Create a new text file
    >>> textfile = make_textfile()
    >>> # Copy the textfile
    >>> copy_file(textfile)
    C:\\Users\\<username>\\generated_textfile.txt'

    Keywords
        make textfile, textfile, testfile, exampel file, make file, make, new file, new textfile, txt, new txt

    Icon
        las la-copy
    
    """
    from uuid import uuid4
    import os
    import shutil

    if not new_path:
        new_path = old_path

    if os.path.isfile(old_path):
        if not os.path.isfile(new_path):
            shutil.copy(old_path, new_path)
        elif os.path.isfile(new_path):
            filename, file_extension = os.path.splitext(old_path)
            new_path = filename + "_copy_" + str(uuid4())[:4] + file_extension
        shutil.copy(old_path, new_path)
    return new_path


@activity
def get_file_extension(path):
    """Get file extension
    
    Get extension of a file

    :parameter path: Path to file to get extension from

    :return: String with extension, e.g. '.txt'

        :Example:

    >>> # Create a new text file
    >>> textfile = make_textfile()
    >>> # Get file extension of this textfile
    >>> get_file_extension(textfile)
    '.txt'

    Keywords
        file, extension, file extension, details

    Icon
        las la-info

    """

    import os
    filename, file_extension = os.path.splitext(path)

    return file_extension


@activity
def send_to_printer(file):
    """Print
    
    Send file to default printer to priner. This activity sends a file to the printer. Make sure to have a default printer set up.

    :parameter file: Path to the file to print, should be a printable file

        :Example:

    >>> # Caution as this example could result in a print from default printer
    >>> # Create a new text file
    >>> textfile = make_textfile(text = 'What does a robot do at lunch? Take a megabyte!')
    >>> # Print the textfile
    >>> send_to_printer(textfile)

    Keywords
        print, printer, printing, ink, export

    Icon
        las la-print
    """
    import os
    os.startfile(file, 'print')


"""
PDF
Icon: las la-file-pdf
"""

@activity
def read_text_from_pdf(file_path):
    """Text from PDF
    
    Extracts the text from a PDF. This activity reads text from a pdf file. Can only read PDF files that contain a text layer.

    :parameter file_path: Path to the PDF (either relative or absolute)
    :return: The text from the PDF

        :Example:
        
    >>> # Caution, for this example to work a .pdf example file will be downloaded from automagica.com FTP
    >>> example_pdf = download_file_from_url('http://automagica.com/examples/example_document.pdf')
    >>> # Open example pdf for illustration
    >>> open_file(example_pdf)
    >>> # Read the text
    >>> read_text_from_pdf(example_pdf)

    Keywords
        PDF, read, text, extract text, PDF file

    Icon
        las la-glasses
    """
    from PyPDF2 import PdfFileReader

    text = ""

    with open(file_path, 'rb') as f:
        reader = PdfFileReader(f)
        for i in range(reader.numPages):
            page = reader.getPage(i)
            text += page.extractText()

    return text


@activity
def join_pdf_files(file_paths, output_path=None):
    """Merge PDF
    
    Merges multiple PDFs into a single file

    :parameter file_paths: List of paths to PDF files
    :parameter output_path: Full path where joined pdf files can be written. If no path is given will write to home dir as 'merged_pdf.pdf'

    :return: Output path as string
    
        :Example:
        
    >>> # Caution, for this example to work a .pdf example file will be downloaded from automagica.com FTP
    >>> example_pdf = download_file_from_url('http://automagica.com/examples/example_document.pdf')
    >>> # Join the PDF file three times with itself for illustration, could also be different files
    >>> merged_pdf = join_pdf_files([example_pdf, example_pdf, example_pdf])
    >>> # Open resulting PDF file for illustration
    >>> open_file(merged_pdf)

    Keywords
        PDF, read, text, extract text, PDF file, join PDF, join, merge, merge PDF

    Icon
        las la-object-ungroup

    """
    from PyPDF2 import PdfFileMerger, PdfFileReader

    if not output_path:
        import os
        output_path = os.path.expanduser("~") + '\merged_pdf.pdf'
        
    merger = PdfFileMerger()
    for file_path in file_paths:
        with open(file_path, "rb") as f:
            merger.append(PdfFileReader(f))

    merger.write(output_path)

    return output_path


@activity
def extract_page_range_from_pdf(file_path, start_page, end_page, output_path=None):
    """Extract page from PDF
    
    Extracts a particular range of a PDF to a separate file.

    :parameter file_path: Path to the PDF (either relative or absolute)
    :parameter start_page: Page number to start from, with 0 being the first page
    :parameter end_page: Page number to end with, with 0 being the first page
    :param output_path: Output path, if no path is provided same path as input will be used with 'extracted' added to the name


        :Example:

    >>> # Caution, for this example to work a .pdf example file will be downloaded from automagica.com FTP
    >>> example_pdf = download_file_from_url('http://automagica.com/examples/example_document.pdf')
    >>> # Join the PDF file three times to create multi page
    >>> multi_page_pdf_example = join_pdf_files([example_pdf, example_pdf, example_pdf])
    >>> # Extract some pages from it
    >>> new_file = extract_page_range_from_pdf(multi_page_pdf_example, 1, 2 )
    >>> # Open resulting PDF file for illustration
    >>> open_file(new_file)

    Keywords
        PDF, read, extract text, PDF file, extract PDF, join, cut, cut PDF, extract pages, extract from pdf, select page, page
    Icon
        las la-cut
    
    """
    from PyPDF2 import PdfFileWriter, PdfFileReader

    if not output_path:
        import os
        base, file_extension = os.path.splitext(file_path)
        output_path = base + '_extracted' + file_extension


    with open(file_path, "rb") as f:

        reader = PdfFileReader(f)
        writer = PdfFileWriter()

        for i in range(start_page, end_page):
            writer.addPage(reader.getPage(i))

        with open(output_path, "wb") as f:
            writer.write(f)

    return output_path


@activity
def extract_images_from_pdf(file_path):
    """Extract images from PDF
    
    Save a specific page from a PDF as an image

    :parameter file_path: Full path to store extracted images

        :Example:
        
    >>> # Caution, for this example to work a .pdf example file will be downloaded from automagica.com FTP
    >>> example_pdf = download_file_from_url('http://automagica.com/examples/example_document.pdf')
    >>> # Extract the images
    >>> extract_images_from_pdf(example_pdf)

    Keywords
        PDF, extract images, images, extract text, PDF file, image

    Icon
        las la-icons
        
    """
    from PyPDF2 import PdfFileReader
    from PIL import Image

    extracted_images = []
    with open(file_path, "rb") as f:
        reader = PdfFileReader(f)
        for i in range(reader.getNumPages()):
            page = reader.getPage(i)
            objects = page["/Resources"]["/XObject"].getObject()

            for obj in objects:
                if objects[obj]["/Subtype"] == "/Image":
                    size = (objects[obj]["/Width"], objects[obj]["/Height"])
                    data = objects[obj].getData()

                    if objects[obj]["/ColorSpace"] == "/DeviceRGB":
                        mode = "RGB"
                    else:
                        mode = "P"

                    if objects[obj]["/Filter"] == "/FlateDecode":
                        img = Image.frombytes(mode, size, data)
                        img.save(obj[1:] + ".png")
                        extraced_images.append(obj[1:] + ".png")

                    elif objects[obj]["/Filter"] == "/JPXDecode":
                        img = open(obj[1:] + ".jp2", "wb")
                        extraced_images.append(obj[1:] + ".jp2")
                        img.write(data)
                        img.close()


@activity
def apply_watermark_to_pdf(file_path, watermark_path, output_path=''):
    """Watermark a PDF

    :parameter file_path: Filepath to the document that will be watermarked. Should be pdf file.
    :parameter watermark_path: Filepath to the watermark. Should be pdf file.
    :parameter output_path: Path to save watermarked PDF. If no path is provided same path as input will be used with 'watermarked' added to the name

    :return: Output path as a string
        :Example:
        
    >>> # Caution, for this example to work a .pdf example file will be downloaded from automagica.com FTP
    >>> example_pdf = download_file_from_url('http://automagica.com/examples/example_document.pdf')
    >>> # Download the watermark
    >>> example_watermark = download_file_from_url('http://automagica.com/examples/approved_stamp.pdf')
    >>> # Apply the watermark
    >>> watermarked_file = apply_watermark_to_pdf(example_pdf, example_watermark)
    >>> # Open the file for illustration
    >>> open_file(watermarked_file)

    Keywords
        PDF, extract images, images, extract text, PDF file, image

    Icon
        las la-stamp
    """
    from PyPDF2 import PdfFileWriter, PdfFileReader
    import os

    if not output_path:
        base, file_extension = os.path.splitext(file_path)
        output_path = base + '_watermarked_' + file_extension


    watermark = PdfFileReader(open(watermark_path, "rb"))

    input_file = PdfFileReader(open(file_path, "rb"))

    page_count = input_file.getNumPages()

    output_file = PdfFileWriter()

    for page_number in range(page_count):
        input_page = input_file.getPage(page_number)
        input_page.mergePage(watermark.getPage(0))
        output_file.addPage(input_page)


    with open(output_path, "wb") as outputStream:
        output_file.write(outputStream)

    return output_path


"""
System Monitoring
Icon: las la-wave-square
"""


@activity
def get_cpu_load(measure_time=1):
    """CPU load

    Get average CPU load for all cores.

    :parameter measure_time: Time (seconds) to measure load. Standard measure_time is 1 second.

    :return: Displayed load is an average over measured_time.

        :Example:

    >>> get_cpu_load()
    10.1

    Keywords
        cpu, load, cpuload

    Icon
        las la-microchip

    """
    import psutil

    cpu_measurements = []
    for _ in range(measure_time):
        cpu_measurements.append(psutil.cpu_percent(interval=1))
    return sum(cpu_measurements) / len(cpu_measurements)


@activity
def get_number_of_cpu(logical=True):
    """Count CPU

    Get the number of CPU's in the current system. 

    :parameter logical: Determines if only logical units are added to the count, default value is True.

    :return: Number of CPU Integer

        :Example:

    >>> get_number_of_cpu()
    2

    Keywords
        cpu, count, number of cpu

    Icon
        las la-calculator

    """
    import psutil

    return psutil.cpu_count(logical=logical)


@activity
def get_cpu_frequency():
    """CPU frequency

    Get frequency at which CPU currently operates.

    :return: minimum and maximum frequency

        :Example:

    >>> get_cpu_frequency()
    scpufreq(current=3600.0, min=0.0, max=3600.0)

    Keywords
        cpu, load, cpu frequency

    Icon
        las la-wave-square

    """
    import psutil

    return psutil.cpu_freq()


@activity
def get_cpu_stats():
    """CPU Stats

    Get CPU statistics

    :return: Number of CTX switches, intterupts, soft-interrupts and systemcalls.

        :Example:

    >>> get_cpu_stats()
    scpustats(ctx_switches=735743826, interrupts=1540483897, soft_interrupts=0, syscalls=2060595131)

    Keywords
        cpu, load, cpu frequency, stats, cpu statistics

    Icon
        las la-server
    """
    import psutil

    return psutil.cpu_stats()


@activity
def get_memory_stats(mem_type="swap"):
    """Memory statistics

    Get  memory statistics

    :parameter mem_type: Choose mem_type = 'virtual' for virtual memory, and mem_type = 'swap' for swap memory (standard).

    :return: Total, used, free and percentage in use.

        :Example:

    >>> get_memory_stats()
    sswap(total=24640016384, used=18120818688, free=6519197696, percent=73.5, sin=0, sout=0)

    Keywords
        memory, statistics, usage, ram

    Icon
        las la-memory

    """
    import psutil

    if mem_type == "virtual":
        return psutil.virtual_memory()
    else:
        return psutil.swap_memory()


@activity
def get_disk_stats():
    """Disk stats

    Get disk statistics of main disk

    :return: Total, used, free and percentage in use.

        :Example:

    >>> get_disk_stats()
    sdiskusage(total=999559262208, used=748696350720, free=250862911488, percent=74.9)

    Keywords
        disk usage, disk stats, disk, harddisk, space

    Icon
        las la-save
    """
    import psutil

    return psutil.disk_usage("/")


@activity
def get_disk_partitions():
    """Partition info

    Get disk partition info

    :return: tuple with info for every partition.

        :Example:

    >>> get_disk_paritions()
    [sdiskpart(device='C:\\', mountpoint='C:\\', fstype='NTFS', opts='rw,fixed')]

    Keywords
        disk usage, disk stats, disk, harddisk, space

    Icon
        las la-save

    """
    import psutil

    return psutil.disk_partitions()


@activity
def get_boot_time():
    """Boot time

    Get most recent boot time

    :return: time PC was booted in seconds after the epoch.

        :Example:

    >>> get_boot_time()
    123456789.0

    Keywords
        boot, boot time, boottime, startup, timer

    Icon
        lar la-clock
    """
    import psutil

    return psutil.boot_time()


@activity
def get_time_since_last_boot():
    """Uptime

    Get uptime since last boot

    :return: time since last boot in seconds.

        :Example:

    >>> get_time_since_last_boot()
    1337.0

    Keywords
        boot, boot time, boottime, startup, timer

    Icon
        lar la-clock
    """
    import time
    import psutil

    return time.time() - psutil.boot_time()


"""
Image Processing
Icon: las la-photo-video
"""


@activity
def show_image(path):
    """Show image

    Displays an image specified by the path variable on the default imaging program.

    :parameter path: Full path to image

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, show image, reveal, open image, open

    Icon
        las la-images

    """
    from PIL import Image

    im = Image.open(path)

    return im.show()


@activity
def rotate_image(path, angle=90):
    """Rotate image

    Rotate an image

    :parameter angle: Degrees to rotate image. Note that angles other than 90, 180, 270, 360 can resize the picture. 

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Rotate the image
    >>> rotate_image(testimage)
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, rotate image, 90 degrees, image manipulation, photoshop, paint

    Icon
        las la-undo

    """
    from PIL import Image

    im = Image.open(path)

    return im.rotate(angle, expand=True).save(path)


@activity
def resize_image(path, size):
    """Resize image

    Resizes the image specified by the path variable. 

    :parameter path: Path to the image
    :parameter size:  Tuple with the width and height in pixels. E.g.  (300, 400) gives the image a width of 300 pixels and a height of 400 pixels.

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Resize the image
    >>> resize_image(testimage, size=(100,100))
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, resize image, resize, size, image manipulation, photoshop, paint

    Icon
        las la-expand-arrows-alt
    """
    from PIL import Image

    im = Image.open(path)

    return im.resize(size).save(path)


@activity
def get_image_width(path):
    """Get image width

    Get with of image

    :parameter path: Path to image

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # get image height
    >>> get_image_width(testimage)
    1000

    Keywords
        image, height, width, image height, image width

    Icon
        las la-expand-arrows-alt
    """
    from PIL import Image

    im = Image.open(path)

    width, _ = im.size

    return width


@activity
def get_image_height(path):
    """Get image height

    Get height of image

    :parameter path: Path to image

    :return: Height of image

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # get image height
    >>> get_image_height(testimage)
    1000

    Keywords
        image, height, width, image height, image width

    Icon
        las la-arrows-alt-v

    """
    from PIL import Image

    im = Image.open(path)

    _, height = im.size

    return height


@activity
def crop_image(path, box=None):
    """Crop image

        Crops the image specified by path to a region determined by the box variable.

    :parameter path: Path to image
    :parameter box:  A tuple that defines the left, upper, right and lower pixel coördinate e.g.: (left, upper, right, lower)

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Crop the image
    >>> crop_image(testimage, box = (10,10,100,100))
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, crop, crop image

    Icon
        las la-crop
    """
    from PIL import Image
    im = Image.open(path)
    return im.crop(box).save(path)


@activity
def mirror_image_horizontally(path):
    """Mirror image horizontally

    Mirrors an image with a given path horizontally from left to right.

    :parameter path: Path to image

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Mirror image horizontally
    >>> mirror_image_horizontally(testimage)
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, flip, flip image, mirror, mirror image, horizon, horizontally

    Icon
        las la-caret-up
    """
    from PIL import Image
    im = Image.open(path)
    return im.transpose(Image.FLIP_LEFT_RIGHT).save(path)


@activity
def mirror_image_vertically(path):
    """Mirror image vertically

    Mirrors an image with a given path vertically from top to bottom.

    :parameter path: Path to image

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Mirror image vertically
    >>> mirror_image_vertically(testimage)
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, flip, flip image, mirror, mirror image, vertical, vertically

    Icon
        las la-caret-right
    """
    from PIL import Image
    im = Image.open(path)
    return im.transpose(Image.FLIP_TOP_BOTTOM).save(path)


"""
Process
Icon: las la-play
"""


@activity
def run_manual(task):
    """Windows run

    Use Windows Run to boot a process
    Note this uses keyboard inputs which means this process can be disrupted by interfering inputs

    :parameter task: Name of the task to run e.g. 'mspaint.exe'

        :Example:

    >>> # Open paint with Windows run
    >>> run_manual('mspaint.exe')
    >>> # Open home directory with Windows run
    >>> run_manual(home_path())

    Keywords
        run, open, task, win r, windows run, shell, cmd

    Icon
        las la-cog
    """

    import time

    press_key_combination('win', 'r')
    time.sleep(0.5)

    import platform

    # Set keyboard layout for Windows platform
    if platform.system() == "Windows":
        from win32api import LoadKeyboardLayout
        LoadKeyboardLayout("00000409", 1)

    type_text(task)
    press_key('enter')


@activity
def run(process):
    """Run process

    Use subprocess to open a windows process

    :parameter process: Process to open e.g: 'calc.exe', 'notepad.exe', 'control.exe', 'mspaint.exe'.

        :Example:

    >>> # Open paint with Windows run
    >>> run('mspaint.exe')

    Keywords
        run, open, task, win r, windows run, shell, cmd

    Icon
        las la-play
    """
    import subprocess
    subprocess.Popen(process)


@activity
def is_process_running(name):
    """Check if process is running

    Check if process is running. Validates if given process name (name) is currently running on the system.

    :parameter name: Name of process

    :return: Boolean

        :Example:

    >>> # Open paint with Windows run
    >>> run('mspaint.exe')
    >>> # Check if paint is running
    >>> is_process_running('mspaint.exe')
    True

    Keywords
        run, open, task, win r, windows run, shell, cmd

    Icon
        las la-cogs
    """
    import psutil

    if name:
        for p in psutil.process_iter():
            if name in p.name():
                return True

    return False


@activity
def get_running_processes():
    """Get running processes

    Get names of unique processes currently running on the system.

    :return: List of unique running processes

         :Example:

    >>> # Show all running processes
    >>> get_running_processes()
    ['cmd.exe', 'chrome.exe', ... ]

    Keywords
        process, processes, list processes, running, running processes

    Icon
        las la-list
    """
    import psutil

    process_list = []

    for p in psutil.process_iter():
        process_list.append(p.name())

    return list(set(process_list))


@activity
def kill_process(name=None):
    """Kill process

    Kills a process forcefully

    :parameter name: Name of the process

        :Example:

    >>> # Open paint with Windows run
    >>> run('mspaint.exe')
    >>> # Force paint to close
    >>> kill_process('mspaint.exe')


    Keywords
        run, open, task, win r, windows run, shell, cmd, kill, stop, kill process, stop process, quit, exit

    Icon
        las la-window-close
    """
    import os
    return os.system("taskkill /f /im " + name + " >nul 2>&1")


"""
Optical Character Recognition (OCR)
Icon: las la-glasses
"""


@activity
def extract_text_ocr(path=None):
    """Get text with OCR

    This activity extracts all text from the current screen or an image if a path is specified.

    :parameter path: Path to image from where text will be extracted. If no path is specified a screenshot of current screen will be used.

    :return: String with all text from current screen

        :Example:

    >>> # Make a textfile with some text to recognize
    >>> testfile = make_textfile(text='OCR Example')
    >>> # Open the textfile
    >>> open_file(testfile)
    >>> # Find the text with OCR
    >>> extracted_text = find_text_on_screen_ocr(text='OCR Example')
    >>> # Check if the extracted_text contains the original word
    >>> 'OCR Example' in extracted_text
    True

    Keywords
        OCR, vision, AI, screen, citrix, read, optical character recognition

    Icon
        lab la-readme
    """

    import requests
    import base64
    import os
    import json

    if not path:
        import PIL.ImageGrab
        img = PIL.ImageGrab.grab()
        path = os.path.join(os.path.expanduser("~"), "ocr_temp.jpg")
        img.save(path, "JPEG")

    # Open file and encode as Base 64
    with open(path, 'rb') as f:
        image_base64 = base64.b64encode(f.read()).decode('utf-8')

    # Get Bot API_key
    config_path = os.path.join(os.path.expanduser("~"), "automagica.json")

    # Read JSON
    with open(config_path) as json_file:
        local_data = json.load(json_file)
        api_key = str(local_data['bot_secret'])  # Your API key

    # Prepare data for request
    data = {
        "image_base64": image_base64,
        "api_key": api_key
    }

    # Post request to API
    r = requests.post('https://ocr.automagica.com/', json=data)

    # Print results
    return r.json()['text']


@activity
def find_text_on_screen_ocr(text, criteria=None):
    """Find text on screen with OCR

    This activity finds position (coordinates) of specified text on the current screen using OCR.

    :parameter text: Text to find. Only exact matches are returned.
    :parameter criteria: Criteria to select on if multiple matches are found. If no criteria is specified all matches will be returned. Options are 'first', which returns the first match closest to the upper left corner, 'last' returns the last match closest to the lower right corner, random selects a random match.

    :return: Dictionary or list of dictionaries with matches with following elements: 'h' height in pixels, 'text' the matched text,'w' the width in pixels, 'x' absolute x-coördinate , 'y' absolute y-coördinate. Returns nothing if no matches are found

        :Example:

    >>> # Make a textfile with some text to recognize
    >>> testfile = make_textfile(text='OCR Example')
    >>> # Open the textfile
    >>> open_file(testfile)
    >>> # Find the text with OCR
    >>> find_text_on_screen_ocr(text='OCR Example')

    Keywords
        OCR, vision, AI, screen, citrix, read, optical character recognition

    Icon
        las la-glasses

    """

    import requests
    import base64
    import os
    import json

    import PIL.ImageGrab
    img = PIL.ImageGrab.grab()
    path = os.path.join(os.path.expanduser("~"), "ocr_capture.jpg")
    img.save(path, "JPEG")

    # Open file and encode as Base 64
    with open(path, 'rb') as f:
        image_base64 = base64.b64encode(f.read()).decode('utf-8')

    # Get Bot API_key
    config_path = os.path.join(os.path.expanduser("~"), "automagica.json")

    # Read JSON
    with open(config_path) as json_file:
        local_data = json.load(json_file)
        api_key = str(local_data['bot_secret'])  # Your API key

    # Prepare data for request
    data = {
        "image_base64": image_base64,
        "api_key": api_key
    }

    # Post request to API
    r = requests.post('https://ocr.automagica.com/', json=data)

    # Print results
    data = r.json()['locations']

    # Find all matches
    matches = []
    for item in data:
        if item['text'].lower() == text.lower():
            matches.append(item)

    if not matches:
        return None

    if criteria:
        if len(matches) > 0:
            if criteria == 'first':
                best_match = matches[0]
            if criteria == 'last':
                best_match = matches[-1]
            if criteria == 'random':
                import random
                best_match = random.choice(matches)

            return best_match

    else:
        return matches


@activity
def click_on_text_ocr(text):
    """Click on text with OCR

    This activity clicks on position (coordinates) of specified text on the current screen using OCR.

    :parameter text: Text to find. Only exact matches are returned.

        :Example:

    >>> # Make a textfile with some text to recognize
    >>> testfile = make_textfile(text='OCR Example')
    >>> # Open the textfile
    >>> open_file(testfile)
    >>> # Find the text with OCR and click on it
    >>> click_on_text(text='OCR Example')

    Keywords
        OCR, vision, AI, screen, citrix, read, optical character recognition, click

    Icon
        las la-mouse-pointer
    """
    position = find_text_on_screen_ocr(text, criteria='first')
    if position:
        from pyautogui import click
        x = int(position['x'] + position['w']/2)
        y = int(position['y'] + position['h']/2)
        return click(x=x, y=y)


@activity
def double_click_on_text_ocr(text):
    """Double click on text with OCR

    This activity double clicks on position (coordinates) of specified text on the current screen using OCR.

    :parameter text: Text to find. Only exact matches are returned.

        :Example:

    >>> # Make a textfile with some text to recognize
    >>> testfile = make_textfile(text='OCR Example')
    >>> # Open the textfile
    >>> open_file(testfile)
    >>> # Find the text with OCR and double click on it
    >>> double_click_on_text(text='OCR Example')

    Keywords
        OCR, vision, AI, screen, citrix, read, optical character recognition, click, double click

    Icon
        las la-mouse-pointer

    """

    position = find_text_on_screen_ocr(text, criteria='first')
    if position:
        from pyautogui import doubleClick
        x = int(position['x'] + position['w']/2)
        y = int(position['y'] + position['h']/2)
        return doubleClick(x=x, y=y)


@activity
def right_click_on_text_ocr(text):
    """Right click on text with OCR

    This activity Right clicks on position (coordinates) of specified text on the current screen using OCR.

    :parameter text: Text to find. Only exact matches are returned.

        :Example:

    >>> # Make a textfile with some text to recognize
    >>> testfile = make_textfile(text='OCR Example')
    >>> # Open the textfile
    >>> open_file(testfile)
    >>> # Find the text with OCR and right click on it
    >>> right_click_on_text(text='OCR Example')

    Keywords
        OCR, vision, AI, screen, citrix, read, optical character recognition, click, right click

    Icon
        las la-mouse-pointer
    """
    position = find_text_on_screen_ocr(text, criteria='first')
    if position:
        from pyautogui import rightClick
        x = int(position['x'] + position['w']/2)
        y = int(position['y'] + position['h']/2)
        return rightClick(x=x, y=y)
